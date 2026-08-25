"""PRISMA 2020 27-item checklist with PDF/DOCX/YAML/Markdown export.

Implements the official PRISMA 2020 checklist (Page MJ, et al. *The PRISMA
2020 statement: an updated guideline for reporting systematic reviews.*
BMJ 2021;372:n71) along with the six official extension checklists
(IPD, NMA, ScR, Harms, Abstract, Diagnostic).

Each checklist item is a :class:`PRISMAItem` dataclass; the checklist
itself is a :class:`PRISMAChecklist` container that supports
``to_pdf`` / ``to_docx`` / ``to_yaml`` / ``to_markdown`` / ``to_dict``.

Heavy dependencies (``reportlab``, ``python-docx``, ``PyYAML``) are imported
lazily inside the methods so the module itself is always importable.
"""
#
# MIT License — Academic Research Suite
# Copyright (c) 2026 — see /LICENSE for full text.
#
from __future__ import annotations

import logging
import os
from dataclasses import asdict, dataclass, field
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# PRISMAItem dataclass.
# ---------------------------------------------------------------------------
@dataclass
class PRISMAItem:
    """One row in the PRISMA 2020 checklist.

    Attributes:
        id: Sequential item identifier (1-27 for the standard 2020 checklist).
        section: Section heading (e.g., ``"Title"``, ``"Methods"``,
            ``"Results"``, ``"Discussion"``).
        item_text: Short title / prompt for the item (e.g.,
            ``"Eligibility criteria"``).
        location_in_report: Page/paragraph/table reference where the item
            is addressed in the user's manuscript.
        reported: ``True`` if the user has confirmed this item is addressed
            in the manuscript; ``False`` otherwise.
        notes: Free-text notes (e.g., deviations, clarifications).
    """

    id: int
    section: str
    item_text: str
    location_in_report: str = ""
    reported: bool = False
    notes: str = ""

    def to_dict(self) -> Dict[str, Any]:
        """Serialise to a plain dict (JSON-compatible)."""
        return asdict(self)

    @classmethod
    def from_dict(cls, d: Dict[str, Any]) -> "PRISMAItem":
        """Reconstruct from a dict (tolerant of missing keys)."""
        return cls(
            id=int(d.get("id", 0)),
            section=str(d.get("section", "")),
            item_text=str(d.get("item_text", "")),
            location_in_report=str(d.get("location_in_report", "")),
            reported=bool(d.get("reported", False)),
            notes=str(d.get("notes", "")),
        )


# ---------------------------------------------------------------------------
# The official PRISMA 2020 27 items.
# ---------------------------------------------------------------------------
_DEFAULT_2020_ITEMS: List[PRISMAItem] = [
    PRISMAItem(1,  "Title",        "Identify the report as a systematic review."),
    PRISMAItem(2,  "Abstract",     "See PRISMA 2020 for Abstracts checklist."),
    PRISMAItem(3,  "Introduction", "Describe the rationale for the review in the context of existing knowledge."),
    PRISMAItem(4,  "Introduction", "Provide an explicit statement of the objective(s) or question(s) the review addresses."),
    PRISMAItem(5,  "Methods",      "Indicate whether a review protocol exists and where it can be accessed (registration number, URL)."),
    PRISMAItem(6,  "Methods",      "Specify the inclusion and exclusion criteria with justification."),
    PRISMAItem(7,  "Methods",      "Specify all information sources (e.g. databases, registers) and the date of last search."),
    PRISMAItem(8,  "Methods",      "Present the full search strategies for all databases, registers, and websites, including any filters and limits."),
    PRISMAItem(9,  "Methods",      "Specify the methods used to decide whether a study was eligible for inclusion, including how many reviewers screened each record and whether processes were completed independently."),
    PRISMAItem(10, "Methods",      "Specify the methods used to collect data from reports, including how many reviewers collected data from each report and whether processes were completed independently."),
    PRISMAItem(11, "Methods",      "List and define all outcomes for which data were sought. Specify whether and how data were sought for individual studies or for summary data."),
    PRISMAItem(12, "Methods",      "Specify the methods used to assess risk of bias in the included studies, including the tool(s) used and how the tool(s) were implemented."),
    PRISMAItem(13, "Methods",      "Specify for each outcome the effect measure(s) (e.g. risk ratio, mean difference) used in the synthesis."),
    PRISMAItem(14, "Methods",      "Describe the processes used to decide which studies were eligible for each synthesis, the order of data presentation, and any methods used to combine results of studies."),
    PRISMAItem(15, "Methods",      "Describe any methods used to assess risk of bias due to missing results in a synthesis (arising from reporting biases)."),
    PRISMAItem(16, "Methods",      "Describe any methods used to assess certainty (or confidence) in the body of evidence for an outcome."),
    PRISMAItem(17, "Results",      "Describe the results of the search and selection process, ideally with a flow diagram."),
    PRISMAItem(18, "Results",      "Cite studies that provided data. Present the key characteristics of the included studies (e.g. study design, participants, exposure/intervention, outcomes)."),
    PRISMAItem(19, "Results",      "Present assessments of risk of bias for each included study."),
    PRISMAItem(20, "Results",      "For all outcomes, present for each study: summary statistics for each group and an effect estimate with its precision."),
    PRISMAItem(21, "Results",      "For each synthesis, briefly summarise the characteristics and risk of bias among the studies contributing to that synthesis."),
    PRISMAItem(22, "Results",      "Present assessments of risk of bias due to missing results in a synthesis."),
    PRISMAItem(23, "Results",      "Present assessments of certainty (or confidence) in the body of evidence for each outcome assessed."),
    PRISMAItem(24, "Discussion",   "Provide a general interpretation of the results in the context of other evidence."),
    PRISMAItem(25, "Discussion",   "Discuss any limitations of the evidence included in the review."),
    PRISMAItem(26, "Discussion",   "Discuss any limitations of the review processes used."),
    PRISMAItem(27, "Discussion",   "Provide a general interpretation of the results and important implications for future research and practice. Note registration and protocol availability."),
]


# ---------------------------------------------------------------------------
# PRISMAChecklist class.
# ---------------------------------------------------------------------------
class PRISMAChecklist:
    """Container for the 27-item PRISMA 2020 checklist.

    Provides round-trip serialisation (``to_dict`` / ``from_dict``,
    ``to_yaml`` / ``from_yaml``), a Markdown summary (``to_markdown``),
    publication-grade PDF/DOCX exports (``to_pdf``, ``to_docx``), and
    completion metrics (``completion_rate``, ``missing_items``).

    Example:
        >>> checklist = PRISMAChecklist()
        >>> checklist.items[0].reported = True
        >>> checklist.completion_rate()
        0.037...
    """

    def __init__(self, items: Optional[List[PRISMAItem]] = None) -> None:
        """Initialise the checklist.

        Args:
            items: Optional list of :class:`PRISMAItem` objects. If
                ``None``, loads the default 27 official PRISMA 2020 items.
        """
        self.items: List[PRISMAItem] = (
            list(items) if items is not None else self.default_2020_items()
        )
        logger.debug("PRISMAChecklist initialised with %d items", len(self.items))

    @staticmethod
    def default_2020_items() -> List[PRISMAItem]:
        """Return a fresh copy of the 27 official PRISMA 2020 items."""
        # Deep-copy via from_dict round-trip so callers can mutate freely.
        return [PRISMAItem.from_dict(i.to_dict()) for i in _DEFAULT_2020_ITEMS]

    # ------------------------------------------------------------------
    # Metrics.
    # ------------------------------------------------------------------
    def completion_rate(self) -> float:
        """Return fraction of items marked as ``reported=True``.

        Returns:
            Float in ``[0.0, 1.0]``. Returns ``0.0`` if the checklist is empty.
        """
        if not self.items:
            return 0.0
        return sum(1 for i in self.items if i.reported) / len(self.items)

    def missing_items(self) -> List[PRISMAItem]:
        """Return the subset of items not yet marked as reported."""
        return [i for i in self.items if not i.reported]

    # ------------------------------------------------------------------
    # Serialisation.
    # ------------------------------------------------------------------
    def to_dict(self) -> Dict[str, Any]:
        """Serialise to a JSON-friendly dict."""
        return {
            "checklist": "PRISMA 2020",
            "n_items": len(self.items),
            "completion_rate": round(self.completion_rate(), 4),
            "items": [i.to_dict() for i in self.items],
        }

    @classmethod
    def from_dict(cls, d: Dict[str, Any]) -> "PRISMAChecklist":
        """Reconstruct from a dict produced by :meth:`to_dict`."""
        items = [PRISMAItem.from_dict(i) for i in d.get("items", [])]
        return cls(items=items)

    def to_yaml(self, path: str) -> str:
        """Write the checklist to a YAML file.

        Args:
            path: Output ``.yaml`` path.

        Returns:
            Absolute path written.
        """
        try:
            import yaml  # noqa: WPS433
        except ImportError as exc:  # pragma: no cover - PyYAML always available
            raise RuntimeError("PyYAML is required for to_yaml()") from exc
        out = os.path.abspath(path)
        os.makedirs(os.path.dirname(out) or ".", exist_ok=True)
        with open(out, "w", encoding="utf-8") as f:
            yaml.safe_dump(self.to_dict(), f, allow_unicode=True, sort_keys=False)
        logger.info("PRISMA checklist YAML written to %s", out)
        return out

    @classmethod
    def from_yaml(cls, path: str) -> "PRISMAChecklist":
        """Load a checklist from a YAML file."""
        import yaml  # noqa: WPS433
        with open(path, "r", encoding="utf-8") as f:
            d = yaml.safe_load(f)
        return cls.from_dict(d)

    def to_markdown(self) -> str:
        """Return a Markdown rendering of the checklist.

        The table has columns: ID, Section, Item, Location, Reported, Notes.
        """
        lines = [
            "# PRISMA 2020 Checklist",
            "",
            f"**Items:** {len(self.items)} | "
            f"**Reported:** {sum(1 for i in self.items if i.reported)} | "
            f"**Completion:** {self.completion_rate()*100:.1f}%",
            "",
            "| ID | Section | Item | Location in report | Reported | Notes |",
            "|----|---------|------|--------------------|----------|-------|",
        ]
        for it in self.items:
            rep = "✓" if it.reported else "—"
            loc = _md_escape(it.location_in_report) or "—"
            notes = _md_escape(it.notes) or "—"
            lines.append(
                f"| {it.id} | {_md_escape(it.section)} | "
                f"{_md_escape(it.item_text)} | {loc} | {rep} | {notes} |"
            )
        return "\n".join(lines) + "\n"

    # ------------------------------------------------------------------
    # PDF export (reportlab).
    # ------------------------------------------------------------------
    def to_pdf(self, path: str, title: str = "PRISMA 2020 Checklist") -> str:
        """Render the checklist as a multi-page PDF table.

        Uses reportlab Platypus with a styled ``Table``; CJK fonts are
        registered on demand via the same machinery as
        :mod:`reporting.pdf_report` (Noto Serif SC body / Noto Sans SC
        headings / DejaVu Sans symbols).

        Args:
            path: Output ``.pdf`` path.
            title: Document title (top of first page).

        Returns:
            Absolute path written.
        """
        try:
            from reportlab.lib import colors  # noqa: WPS433
            from reportlab.lib.pagesizes import A4  # noqa: WPS433
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # noqa: WPS433
            from reportlab.lib.units import cm  # noqa: WPS433
            from reportlab.platypus import (  # noqa: WPS433
                Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
            )
        except ImportError as exc:  # pragma: no cover - reportlab in requirements
            raise RuntimeError("reportlab is required for to_pdf()") from exc

        out = os.path.abspath(path)
        os.makedirs(os.path.dirname(out) or ".", exist_ok=True)

        # Register CJK fonts (idempotent).
        _register_reportlab_fonts()
        body_font = _pick_reportlab_font(
            "NotoSerifSC", "NotoSansSC", "Helvetica",
        )
        head_font = _pick_reportlab_font(
            "NotoSansSC-Bold", "NotoSansSC", "Helvetica-Bold",
        )

        doc = SimpleDocTemplate(
            out, pagesize=A4,
            leftMargin=1.6 * cm, rightMargin=1.6 * cm,
            topMargin=1.6 * cm, bottomMargin=1.6 * cm,
            title=title,
        )
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "ChecklistTitle", parent=styles["Heading1"],
            fontName=head_font, fontSize=18, leading=22,
            spaceBefore=0, spaceAfter=4, textColor=colors.HexColor("#1a3a5c"),
        )
        subtitle_style = ParagraphStyle(
            "ChecklistSubtitle", parent=styles["Normal"],
            fontName=body_font, fontSize=9, leading=12,
            textColor=colors.HexColor("#555555"), spaceAfter=12,
        )
        cell_style = ParagraphStyle(
            "ChecklistCell", parent=styles["Normal"],
            fontName=body_font, fontSize=8, leading=10,
        )
        head_style = ParagraphStyle(
            "ChecklistHead", parent=styles["Normal"],
            fontName=head_font, fontSize=9, leading=11,
            textColor=colors.white,
        )

        story: List[Any] = []
        story.append(Paragraph(title, title_style))
        story.append(Paragraph(
            "Page MJ, et al. The PRISMA 2020 statement: an updated guideline "
            "for reporting systematic reviews. BMJ 2021;372:n71.",
            subtitle_style,
        ))
        story.append(Paragraph(
            f"Completion: {self.completion_rate()*100:.1f}% "
            f"({sum(1 for i in self.items if i.reported)}/{len(self.items)} items reported).",
            subtitle_style,
        ))

        # Table header.
        data: List[List[Any]] = [[
            Paragraph("<b>ID</b>", head_style),
            Paragraph("<b>Section</b>", head_style),
            Paragraph("<b>Item</b>", head_style),
            Paragraph("<b>Location</b>", head_style),
            Paragraph("<b>Reported</b>", head_style),
            Paragraph("<b>Notes</b>", head_style),
        ]]
        for it in self.items:
            data.append([
                Paragraph(str(it.id), cell_style),
                Paragraph(_xml_escape(it.section), cell_style),
                Paragraph(_xml_escape(it.item_text), cell_style),
                Paragraph(_xml_escape(it.location_in_report or "—"), cell_style),
                Paragraph("✓" if it.reported else "—", cell_style),
                Paragraph(_xml_escape(it.notes or "—"), cell_style),
            ])

        table = Table(
            data,
            colWidths=[0.9 * cm, 2.0 * cm, 6.2 * cm, 2.7 * cm, 1.5 * cm, 4.2 * cm],
            repeatRows=1,
        )
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a3a5c")),
            ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
            ("ALIGN",      (0, 0), (-1, -1), "LEFT"),
            ("VALIGN",     (0, 0), (-1, -1), "TOP"),
            ("GRID",       (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                [colors.white, colors.HexColor("#f4f7fb")]),
            ("LEFTPADDING",  (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING",   (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 3),
        ]))
        story.append(table)
        doc.build(story)
        logger.info("PRISMA checklist PDF written to %s", out)
        return out

    # ------------------------------------------------------------------
    # DOCX export (python-docx).
    # ------------------------------------------------------------------
    def to_docx(self, path: str, title: str = "PRISMA 2020 Checklist") -> str:
        """Render the checklist as a Word document.

        Args:
            path: Output ``.docx`` path.
            title: Document title.

        Returns:
            Absolute path written.
        """
        try:
            from docx import Document  # noqa: WPS433
            from docx.shared import Pt, RGBColor, Inches  # noqa: WPS433
            from docx.enum.table import WD_ALIGN_VERTICAL  # noqa: WPS433
            from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: WPS433
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError("python-docx is required for to_docx()") from exc

        out = os.path.abspath(path)
        os.makedirs(os.path.dirname(out) or ".", exist_ok=True)

        doc = Document()
        # Set default font.
        normal = doc.styles["Normal"]
        normal.font.name = "Noto Sans SC"
        normal.font.size = Pt(10)

        # Title.
        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        run = p.add_run(
            "Page MJ, et al. The PRISMA 2020 statement: an updated guideline "
            "for reporting systematic reviews. BMJ 2021;372:n71."
        )
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        summary = doc.add_paragraph()
        summary.add_run(
            f"Completion: {self.completion_rate()*100:.1f}% "
            f"({sum(1 for i in self.items if i.reported)}/{len(self.items)} items reported)."
        )

        # Table.
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, label in enumerate(["ID", "Section", "Item",
                                   "Location", "Reported", "Notes"]):
            hdr[i].text = label
            for para in hdr[i].paragraphs:
                for r in para.runs:
                    r.bold = True

        for it in self.items:
            row = table.add_row().cells
            row[0].text = str(it.id)
            row[1].text = it.section
            row[2].text = it.item_text
            row[3].text = it.location_in_report or "—"
            row[4].text = "✓" if it.reported else "—"
            row[5].text = it.notes or "—"

        doc.save(out)
        logger.info("PRISMA checklist DOCX written to %s", out)
        return out


# ---------------------------------------------------------------------------
# PRISMAExtensionsChecklist class — per-extension item sets.
# ---------------------------------------------------------------------------
class PRISMAExtensionsChecklist:
    """Returns extension-specific checklist item lists.

    Each method returns a ``List[PRISMAItem]`` of additional items that
    should be appended to the standard 27-item PRISMA 2020 checklist when
    reporting a systematic review of the relevant type.
    """

    @staticmethod
    def ipd_checklist() -> List[PRISMAItem]:
        """Return IPD-specific items (Stewart et al., BMJ 2012;345:e5705)."""
        return [
            PRISMAItem(101, "Methods (IPD)",
                       "Describe the methods used to request, receive, and check individual participant data (IPD) from study authors."),
            PRISMAItem(102, "Methods (IPD)",
                       "Specify the methods used to re-analyse IPD (e.g. one-stage vs two-stage approach)."),
            PRISMAItem(103, "Methods (IPD)",
                       "Describe any methods used to assess risk of bias in IPD studies (e.g. updated Cochrane risk-of-bias tool)."),
            PRISMAItem(104, "Methods (IPD)",
                       "Specify the methods used to assess the certainty (confidence) in IPD syntheses."),
            PRISMAItem(105, "Results (IPD)",
                       "Report the number of studies and participants for which IPD were sought and obtained."),
            PRISMAItem(106, "Results (IPD)",
                       "Present IPD synthesis results (e.g. forest plots with IPD-based effect estimates and confidence intervals)."),
        ]

    @staticmethod
    def nma_checklist() -> List[PRISMAItem]:
        """Return NMA-specific items (Hutton et al., Ann Intern Med 2015;162:777)."""
        return [
            PRISMAItem(201, "Methods (NMA)",
                       "Describe the methods used for the network geometry (e.g. network plots), including treatment nodes and edges."),
            PRISMAItem(202, "Methods (NMA)",
                       "Describe the methods used to assess the assumptions of transitivity and consistency."),
            PRISMAItem(203, "Methods (NMA)",
                       "Specify the methods used for the NMA (e.g. Bayesian vs frequentist, choice of model)."),
            PRISMAItem(204, "Methods (NMA)",
                       "Describe methods used to evaluate inconsistency between direct and indirect evidence."),
            PRISMAItem(205, "Results (NMA)",
                       "Present the network geometry of interventions and studies (network plot)."),
            PRISMAItem(206, "Results (NMA)",
                       "Present summary effect estimates from the NMA (e.g. league table, ranking probabilities)."),
        ]

    @staticmethod
    def scr_checklist() -> List[PRISMAItem]:
        """Return Scoping Review items (Tricco et al., Ann Intern Med 2018;169:467)."""
        return [
            PRISMAItem(301, "Methods (ScR)",
                       "Specify the framework or rationale for conducting a scoping review (e.g. Arksey & O'Malley, JBI)."),
            PRISMAItem(302, "Methods (ScR)",
                       "Indicate whether a charting form was developed and used; describe how data were extracted and charted."),
            PRISMAItem(303, "Methods (ScR)",
                       "Note that risk-of-bias assessment was not performed (per scoping review methodology)."),
            PRISMAItem(304, "Methods (ScR)",
                       "Describe any methods used to summarise or synthesise charted data (e.g. narrative, descriptive numerical)."),
            PRISMAItem(305, "Results (ScR)",
                       "Present the charted results (e.g. characteristics of sources, key concepts, gaps identified)."),
            PRISMAItem(306, "Discussion (ScR)",
                       "Describe the implications of the results for practice, policy, and future research."),
        ]

    @staticmethod
    def harms_checklist() -> List[PRISMAItem]:
        """Return adverse events items (Zorzela et al., PLoS One 2016;11:e0157635)."""
        return [
            PRISMAItem(401, "Methods (Harms)",
                       "Specify the methods used to collect adverse events data, including severity grading (e.g. CTCAE) and seriousness."),
            PRISMAItem(402, "Methods (Harms)",
                       "Describe the tools used to assess risk of bias for adverse events outcomes (e.g. revised Cochrane tool, ROBINS-I)."),
            PRISMAItem(403, "Methods (Harms)",
                       "Specify whether and how adverse events were grouped (by system, organ, severity, relatedness)."),
            PRISMAItem(404, "Results (Harms)",
                       "Present the numbers and types of adverse events per study and per arm."),
            PRISMAItem(405, "Results (Harms)",
                       "Present summary effect estimates for adverse events (if meta-analysed)."),
            PRISMAItem(406, "Discussion (Harms)",
                       "Discuss the certainty (confidence) in adverse events evidence."),
        ]

    @staticmethod
    def abstract_checklist() -> List[PRISMAItem]:
        """Return conference abstract items (Beller et al., J Clin Epidemiol 2013;66:657)."""
        return [
            PRISMAItem(501, "Abstract",
                       "Provide the design of the review (e.g. systematic review, meta-analysis)."),
            PRISMAItem(502, "Abstract",
                       "State the objectives of the review using PICO or similar."),
            PRISMAItem(503, "Abstract",
                       "Describe the eligibility criteria and information sources."),
            PRISMAItem(504, "Abstract",
                       "State the methods of synthesis and risk-of-bias assessment."),
            PRISMAItem(505, "Abstract",
                       "Provide the number of included studies and participants."),
            PRISMAItem(506, "Abstract",
                       "State the principal findings and the strength of evidence."),
            PRISMAItem(507, "Abstract",
                       "State funding sources and any conflicts of interest."),
        ]

    @staticmethod
    def diagnostic_checklist() -> List[PRISMAItem]:
        """Return diagnostic test accuracy items (McInnes et al., JAMA 2018;319:388)."""
        return [
            PRISMAItem(601, "Methods (DTA)",
                       "Describe the index test(s) and how they were conducted and interpreted."),
            PRISMAItem(602, "Methods (DTA)",
                       "Describe the reference standard and how it was conducted and interpreted."),
            PRISMAItem(603, "Methods (DTA)",
                       "Specify the methods for collecting or deriving 2×2 data (true positives, false positives, false negatives, true negatives)."),
            PRISMAItem(604, "Methods (DTA)",
                       "Specify the methods used to assess risk of bias and applicability (e.g. QUADAS-2)."),
            PRISMAItem(605, "Methods (DTA)",
                       "Specify the methods of meta-analysis (e.g. bivariate / HSROC model, hierarchical summary ROC)."),
            PRISMAItem(606, "Results (DTA)",
                       "Present the numbers of studies, participants, and test results."),
            PRISMAItem(607, "Results (DTA)",
                       "Present summary estimates of diagnostic accuracy (sensitivity, specificity, likelihood ratios) with confidence intervals."),
        ]


# ---------------------------------------------------------------------------
# ReportLab font registration (mirrors reporting.pdf_report._register_fonts).
# ---------------------------------------------------------------------------
_REPORTLAB_FONTS_REGISTERED = False
_REPORTLAB_AVAILABLE: Dict[str, str] = {}

_RL_FONT_CANDIDATES: Dict[str, List[str]] = {
    "NotoSerifSC": [
        "/usr/share/fonts/truetype/noto-serif-sc/NotoSerifSC-Regular.ttf",
        "/usr/share/fonts/opentype/noto/NotoSerifSC-Regular.otf",
        "/usr/share/fonts/noto-cjk/NotoSerifCJKsc-Regular.otf",
    ],
    "NotoSerifSC-Bold": [
        "/usr/share/fonts/truetype/noto-serif-sc/NotoSerifSC-Bold.ttf",
        "/usr/share/fonts/opentype/noto/NotoSerifSC-Bold.otf",
    ],
    "NotoSansSC": [
        "/usr/share/fonts/truetype/chinese/NotoSansSC-Regular.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf",
        "/usr/share/fonts/noto-cjk/NotoSansCJKsc-Regular.otf",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
    ],
    "NotoSansSC-Bold": [
        "/usr/share/fonts/truetype/chinese/NotoSansSC-Bold.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJKsc-Bold.otf",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
    ],
    "DejaVuSans": [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ],
    "DejaVuSans-Bold": [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    ],
}


def _register_reportlab_fonts() -> None:
    """Register CJK + symbol fonts with reportlab (idempotent)."""
    global _REPORTLAB_FONTS_REGISTERED, _REPORTLAB_AVAILABLE
    if _REPORTLAB_FONTS_REGISTERED:
        return
    _REPORTLAB_FONTS_REGISTERED = True
    try:
        from reportlab.pdfbase import pdfmetrics  # noqa: WPS433
        from reportlab.pdfbase.ttfonts import TTFont  # noqa: WPS433
        from reportlab.pdfbase.pdfmetrics import registerFontFamily  # noqa: WPS433
    except ImportError:  # pragma: no cover - reportlab always in requirements
        return
    for font_name, candidates in _RL_FONT_CANDIDATES.items():
        for path in candidates:
            if not os.path.exists(path):
                continue
            try:
                kwargs: Dict[str, Any] = {}
                if path.endswith(".ttc"):
                    kwargs["subfontIndex"] = 0
                f = TTFont(font_name, path, **kwargs)
                pdfmetrics.registerFont(f)
                _REPORTLAB_AVAILABLE[font_name] = path
                break
            except Exception as exc:
                logger.debug("font register failed %s @ %s: %s", font_name, path, exc)
                continue
    # Register font families so <b>/<i> tags work inside Paragraph.
    def _family(name: str, normal: str, bold: str) -> None:
        if normal in _REPORTLAB_AVAILABLE and bold in _REPORTLAB_AVAILABLE:
            try:
                registerFontFamily(name, normal=normal, bold=bold,
                                   italic=normal, boldItalic=bold)
            except Exception:  # pragma: no cover
                pass
    _family("NotoSerifSC", "NotoSerifSC", "NotoSerifSC-Bold")
    _family("NotoSansSC",  "NotoSansSC",  "NotoSansSC-Bold")
    _family("DejaVuSans",  "DejaVuSans",  "DejaVuSans-Bold")


def _pick_reportlab_font(*candidates: str) -> str:
    """Return the first candidate font that was successfully registered."""
    for c in candidates:
        if c in _REPORTLAB_AVAILABLE:
            return c
    return candidates[-1]


# ---------------------------------------------------------------------------
# Helpers.
# ---------------------------------------------------------------------------
def _md_escape(text: str) -> str:
    """Escape pipe characters in Markdown table cells."""
    if not text:
        return ""
    return text.replace("|", "\\|").replace("\n", " ")


def _xml_escape(text: str) -> str:
    """Escape ``&``, ``<``, ``>`` for safe ReportLab Paragraph inclusion."""
    if not text:
        return ""
    s = str(text)
    s = s.replace("&", "&amp;")
    s = s.replace("<", "&lt;")
    s = s.replace(">", "&gt;")
    return s


__all__ = [
    "PRISMAItem",
    "PRISMAChecklist",
    "PRISMAExtensionsChecklist",
]
