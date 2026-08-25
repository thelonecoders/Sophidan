"""Complete PRISMA report generator — flow + checklist + extraction table.

Combines :class:`prisma.flow_diagram.PRISMAFlowGenerator`,
:class:`prisma.checklist.PRISMAChecklist`, and a list of
:class:`prisma.extraction_form.PRISMAExtractionForm` records into a single
publication-grade report (PDF / DOCX / HTML / Markdown).

The report contains:

1. Title page with review title, authors, and PRISMA 2020 attribution.
2. PRISMA 2020 flow diagram.
3. Completed 27-item PRISMA 2020 checklist.
4. Table of included studies with extracted data.
5. Summary section (counts, completion rate).

Heavy dependencies (``reportlab``, ``python-docx``, ``matplotlib``) are
imported lazily inside the methods so the module itself is always importable.
"""
#
# MIT License — Academic Research Suite
# Copyright (c) 2026 — see /LICENSE for full text.
#
from __future__ import annotations

import io
import logging
import os
import tempfile
from typing import Any, Dict, List, Optional

from .checklist import (
    PRISMAChecklist,
    PRISMAItem,
    _register_reportlab_fonts,
    _pick_reportlab_font,
    _xml_escape,
)
from .extraction_form import PRISMAExtractionForm
from .flow_diagram import PRISMAFlowGenerator, PRISMAStageCounts

logger = logging.getLogger(__name__)


class PRISMAReport:
    """End-to-end PRISMA report generator.

    Bundles a PRISMA 2020 flow diagram, the completed 27-item checklist,
    and a table of included studies with extracted data into a single
    publication-grade document.

    Example:
        >>> from prisma.report import PRISMAReport
        >>> from prisma.flow_diagram import PRISMAStageCounts
        >>> from prisma.checklist import PRISMAChecklist
        >>> from prisma.extraction_form import PRISMAExtractionForm
        >>> counts = PRISMAStageCounts(n_records_databases=500, n_records_screened=400)
        >>> checklist = PRISMAChecklist()
        >>> checklist.items[0].reported = True
        >>> extraction = [PRISMAExtractionForm(study_id='Smith2020', study_design='RCT')]
        >>> report = PRISMAReport(counts, checklist, extraction,
        ...                       title='Example Systematic Review',
        ...                       authors=['Doe J', 'Smith A'])
        >>> report.generate('/tmp/report.pdf', format='pdf')
    """

    def __init__(
        self,
        counts: PRISMAStageCounts,
        checklist: PRISMAChecklist,
        extraction: List[PRISMAExtractionForm],
        title: str = "",
        authors: Optional[List[str]] = None,
    ) -> None:
        """Initialise the report.

        Args:
            counts: Stage counts for the flow diagram.
            checklist: A completed (or in-progress) PRISMA 2020 checklist.
            extraction: List of per-study extraction forms (one per
                included study).
            title: Review title.
            authors: List of author names.
        """
        self.counts = counts
        self.checklist = checklist
        self.extraction = list(extraction) if extraction else []
        self.title = title or "PRISMA 2020 Report"
        self.authors = list(authors) if authors else []
        logger.debug(
            "PRISMAReport created: title=%r, authors=%d, extraction=%d",
            self.title, len(self.authors), len(self.extraction),
        )

    # ------------------------------------------------------------------
    # Top-level generate.
    # ------------------------------------------------------------------
    def generate(self, output_path: str, format: str = "pdf") -> str:
        """Generate the full report in the requested format.

        Args:
            output_path: Where to write the output file.
            format: One of ``'pdf'``, ``'docx'``, ``'html'``.

        Returns:
            Absolute path to the written file.
        """
        format = format.lower().lstrip(".")
        if format == "pdf":
            return self._generate_pdf(output_path)
        if format == "docx":
            return self._generate_docx(output_path)
        if format == "html":
            return self._generate_html(output_path)
        raise ValueError(
            f"format must be one of 'pdf', 'docx', 'html'; got {format!r}"
        )

    def to_markdown(self) -> str:
        """Return the full report as a Markdown string.

        Sections: Title, Authors, Summary counts, Flow-diagram DOT,
        Checklist table, Extraction table, Notes.
        """
        lines: List[str] = []
        lines.append(f"# {self.title}")
        lines.append("")
        if self.authors:
            lines.append("**Authors:** " + ", ".join(self.authors))
            lines.append("")
        lines.append("## Summary")
        lines.append("")
        c = self.counts
        lines.append(f"- Records identified from databases: "
                     f"{c.n_records_databases if c.n_records_databases is not None else '—'}")
        lines.append(f"- Records identified from registers: "
                     f"{c.n_records_registers if c.n_records_registers is not None else '—'}")
        lines.append(f"- Duplicates removed: "
                     f"{c.n_duplicates_removed if c.n_duplicates_removed is not None else '—'}")
        lines.append(f"- Records screened: "
                     f"{c.n_records_screened if c.n_records_screened is not None else '—'}")
        lines.append(f"- Records excluded at title/abstract: "
                     f"{c.n_records_excluded_title_abstract if c.n_records_excluded_title_abstract is not None else '—'}")
        lines.append(f"- Full-text assessed: "
                     f"{c.n_full_text_assessed if c.n_full_text_assessed is not None else '—'}")
        lines.append(f"- Studies in qualitative synthesis: "
                     f"{c.n_studies_included_qualitative if c.n_studies_included_qualitative is not None else '—'}")
        lines.append(f"- Studies in quantitative synthesis: "
                     f"{c.n_studies_included_quantitative if c.n_studies_included_quantitative is not None else '—'}")
        lines.append("")

        # Flow diagram (DOT representation).
        gen = PRISMAFlowGenerator(self.counts, title=self.title, extension="standard")
        lines.append("## PRISMA 2020 Flow Diagram")
        lines.append("")
        lines.append("```dot")
        lines.append(gen.to_dot())
        lines.append("```")
        lines.append("")

        # Checklist.
        lines.append(self.checklist.to_markdown())
        lines.append("")

        # Extraction table.
        lines.append("## Included studies — extraction table")
        lines.append("")
        if not self.extraction:
            lines.append("_No extraction forms provided._")
        else:
            lines.append(
                "| Study ID | Design | Population | Intervention | "
                "Comparator | Sample | Follow-up |"
            )
            lines.append(
                "|----------|--------|------------|---------------|"
                "-----------|--------|-----------|"
            )
            for f in self.extraction:
                n = f.sample_size if f.sample_size is not None else "—"
                lines.append(
                    f"| {f.study_id or '—'} | {f.study_design or '—'} | "
                    f"{f.population or '—'} | {f.intervention or '—'} | "
                    f"{f.comparator or '—'} | {n} | {f.follow_up or '—'} |"
                )
        lines.append("")
        lines.append("---")
        lines.append("")
        lines.append(
            "_Generated by Academic Research Suite v2.0.0 — "
            "PRISMA 2020 (Page MJ et al., BMJ 2021;372:n71)._"
        )
        return "\n".join(lines) + "\n"

    # ------------------------------------------------------------------
    # PDF generation.
    # ------------------------------------------------------------------
    def _generate_pdf(self, output_path: str) -> str:
        """Render the full report as a multi-section PDF (reportlab)."""
        try:
            from reportlab.lib import colors  # noqa: WPS433
            from reportlab.lib.pagesizes import A4  # noqa: WPS433
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # noqa: WPS433
            from reportlab.lib.units import cm  # noqa: WPS433
            from reportlab.platypus import (  # noqa: WPS433
                BaseDocTemplate, Frame, Image, NextPageTemplate,
                PageBreak, PageTemplate, Paragraph, SimpleDocTemplate,
                Spacer, Table, TableStyle, KeepTogether,
            )
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError("reportlab is required for PDF generation") from exc

        out = os.path.abspath(output_path)
        os.makedirs(os.path.dirname(out) or ".", exist_ok=True)
        _register_reportlab_fonts()
        body_font = _pick_reportlab_font(
            "NotoSerifSC", "NotoSansSC", "Helvetica",
        )
        head_font = _pick_reportlab_font(
            "NotoSansSC-Bold", "NotoSansSC", "Helvetica-Bold",
        )

        # Render the flow diagram PNG to embed.
        gen = PRISMAFlowGenerator(self.counts, title=self.title, extension="standard")
        with tempfile.NamedTemporaryFile(
            mode="w+b", suffix=".png", delete=False,
        ) as tmp_fig:
            flow_png = tmp_fig.name
        try:
            gen.render_png(flow_png, dpi=150)
        except Exception as exc:  # pragma: no cover
            logger.warning("Flow PNG render failed: %s", exc)
            flow_png = ""

        doc = SimpleDocTemplate(
            out, pagesize=A4,
            leftMargin=1.8 * cm, rightMargin=1.8 * cm,
            topMargin=2.0 * cm, bottomMargin=2.0 * cm,
            title=self.title,
        )
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "ReportTitle", parent=styles["Title"],
            fontName=head_font, fontSize=22, leading=28,
            textColor=colors.HexColor("#1a3a5c"),
        )
        author_style = ParagraphStyle(
            "ReportAuthors", parent=styles["Normal"],
            fontName=body_font, fontSize=11, leading=14,
            textColor=colors.HexColor("#444444"),
            alignment=1,  # centre
        )
        h1_style = ParagraphStyle(
            "ReportH1", parent=styles["Heading1"],
            fontName=head_font, fontSize=16, leading=20,
            textColor=colors.HexColor("#1a3a5c"),
            spaceBefore=12, spaceAfter=6,
        )
        body_style = ParagraphStyle(
            "ReportBody", parent=styles["Normal"],
            fontName=body_font, fontSize=10, leading=14,
        )
        cell_style = ParagraphStyle(
            "ReportCell", parent=styles["Normal"],
            fontName=body_font, fontSize=8, leading=10,
        )
        head_cell_style = ParagraphStyle(
            "ReportHeadCell", parent=styles["Normal"],
            fontName=head_font, fontSize=9, leading=11,
            textColor=colors.white,
        )
        attribution_style = ParagraphStyle(
            "ReportAttr", parent=styles["Normal"],
            fontName=body_font, fontSize=8, leading=10,
            textColor=colors.HexColor("#666666"), alignment=1,
        )

        story: List[Any] = []

        # ---- Title page ----
        story.append(Spacer(1, 4 * cm))
        story.append(Paragraph(_xml_escape(self.title), title_style))
        if self.authors:
            story.append(Spacer(1, 0.5 * cm))
            story.append(Paragraph(
                _xml_escape(", ".join(self.authors)), author_style,
            ))
        story.append(Spacer(1, 2 * cm))
        story.append(Paragraph(
            "PRISMA 2020 Reporting Package",
            author_style,
        ))
        story.append(Paragraph(
            "Page MJ, et al. <i>The PRISMA 2020 statement: an updated guideline "
            "for reporting systematic reviews.</i> BMJ 2021;372:n71.",
            attribution_style,
        ))
        story.append(PageBreak())

        # ---- PRISMA Flow Diagram page ----
        story.append(Paragraph("PRISMA 2020 Flow Diagram", h1_style))
        if flow_png and os.path.exists(flow_png):
            try:
                story.append(Image(flow_png, width=16 * cm, height=22 * cm,
                                   kind="proportional"))
            except Exception:  # pragma: no cover
                pass
        else:
            story.append(Paragraph(
                "(Flow diagram could not be rendered.)", body_style,
            ))
        story.append(PageBreak())

        # ---- Checklist ----
        story.append(Paragraph("PRISMA 2020 Checklist", h1_style))
        story.append(Paragraph(
            f"Completion: {self.checklist.completion_rate()*100:.1f}% "
            f"({sum(1 for i in self.checklist.items if i.reported)}/"
            f"{len(self.checklist.items)} items reported).",
            body_style,
        ))
        story.append(Spacer(1, 0.4 * cm))

        # Checklist table.
        data: List[List[Any]] = [[
            Paragraph("<b>ID</b>", head_cell_style),
            Paragraph("<b>Section</b>", head_cell_style),
            Paragraph("<b>Item</b>", head_cell_style),
            Paragraph("<b>Location</b>", head_cell_style),
            Paragraph("<b>Reported</b>", head_cell_style),
        ]]
        for it in self.checklist.items:
            data.append([
                Paragraph(str(it.id), cell_style),
                Paragraph(_xml_escape(it.section), cell_style),
                Paragraph(_xml_escape(it.item_text), cell_style),
                Paragraph(_xml_escape(it.location_in_report or "—"), cell_style),
                Paragraph("✓" if it.reported else "—", cell_style),
            ])
        table = Table(
            data,
            colWidths=[0.9 * cm, 2.2 * cm, 8.5 * cm, 3.5 * cm, 1.7 * cm],
            repeatRows=1,
        )
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a3a5c")),
            ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
            ("ALIGN",      (0, 0), (-1, -1), "LEFT"),
            ("VALIGN",     (0, 0), (-1, -1), "TOP"),
            ("GRID",       (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                [colors.white, colors.HexColor("#f4f7fb")]),
            ("LEFTPADDING",  (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING",   (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 3),
        ]))
        story.append(table)
        story.append(PageBreak())

        # ---- Included studies table ----
        story.append(Paragraph("Included Studies", h1_style))
        story.append(Paragraph(
            f"{len(self.extraction)} study extraction form(s) included.",
            body_style,
        ))
        story.append(Spacer(1, 0.4 * cm))

        if not self.extraction:
            story.append(Paragraph(
                "<i>No extraction forms provided.</i>", body_style,
            ))
        else:
            ex_data: List[List[Any]] = [[
                Paragraph("<b>ID</b>", head_cell_style),
                Paragraph("<b>Design</b>", head_cell_style),
                Paragraph("<b>Population</b>", head_cell_style),
                Paragraph("<b>Intervention</b>", head_cell_style),
                Paragraph("<b>Comparator</b>", head_cell_style),
                Paragraph("<b>N</b>", head_cell_style),
            ]]
            for f in self.extraction:
                n = f.sample_size if f.sample_size is not None else "—"
                ex_data.append([
                    Paragraph(_xml_escape(f.study_id or "—"), cell_style),
                    Paragraph(_xml_escape(f.study_design or "—"), cell_style),
                    Paragraph(_xml_escape(f.population or "—"), cell_style),
                    Paragraph(_xml_escape(f.intervention or "—"), cell_style),
                    Paragraph(_xml_escape(f.comparator or "—"), cell_style),
                    Paragraph(str(n), cell_style),
                ])
            ex_table = Table(
                ex_data,
                colWidths=[2.5 * cm, 2.0 * cm, 4.0 * cm, 4.0 * cm,
                           3.0 * cm, 1.3 * cm],
                repeatRows=1,
            )
            ex_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a3a5c")),
                ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
                ("ALIGN",      (0, 0), (-1, -1), "LEFT"),
                ("VALIGN",     (0, 0), (-1, -1), "TOP"),
                ("GRID",       (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                    [colors.white, colors.HexColor("#f4f7fb")]),
            ]))
            story.append(ex_table)

        # ---- Summary ----
        story.append(Spacer(1, 0.6 * cm))
        story.append(Paragraph("Summary", h1_style))
        c = self.counts
        summary_lines = [
            f"<b>Records identified from databases:</b> {c.n_records_databases if c.n_records_databases is not None else '—'}",
            f"<b>Records identified from registers:</b> {c.n_records_registers if c.n_records_registers is not None else '—'}",
            f"<b>Duplicates removed:</b> {c.n_duplicates_removed if c.n_duplicates_removed is not None else '—'}",
            f"<b>Records screened:</b> {c.n_records_screened if c.n_records_screened is not None else '—'}",
            f"<b>Records excluded at title/abstract:</b> {c.n_records_excluded_title_abstract if c.n_records_excluded_title_abstract is not None else '—'}",
            f"<b>Full-text assessed:</b> {c.n_full_text_assessed if c.n_full_text_assessed is not None else '—'}",
            f"<b>Studies in qualitative synthesis:</b> {c.n_studies_included_qualitative if c.n_studies_included_qualitative is not None else '—'}",
            f"<b>Studies in quantitative synthesis:</b> {c.n_studies_included_quantitative if c.n_studies_included_quantitative is not None else '—'}",
            f"<b>Checklist completion:</b> {self.checklist.completion_rate()*100:.1f}%",
            f"<b>Included studies (extraction forms):</b> {len(self.extraction)}",
        ]
        for s in summary_lines:
            story.append(Paragraph(s, body_style))
        story.append(Spacer(1, 1 * cm))
        story.append(Paragraph(
            "Generated by Academic Research Suite v2.0.0 — "
            "PRISMA 2020 (Page MJ et al., BMJ 2021;372:n71).",
            attribution_style,
        ))

        doc.build(story)
        if flow_png and os.path.exists(flow_png):
            try:
                os.unlink(flow_png)
            except OSError:
                pass
        logger.info("PRISMA report PDF written to %s", out)
        return out

    # ------------------------------------------------------------------
    # DOCX generation.
    # ------------------------------------------------------------------
    def _generate_docx(self, output_path: str) -> str:
        """Render the full report as a Word .docx (python-docx)."""
        try:
            from docx import Document  # noqa: WPS433
            from docx.shared import Pt, RGBColor, Inches  # noqa: WPS433
            from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: WPS433
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError("python-docx is required for DOCX generation") from exc

        out = os.path.abspath(output_path)
        os.makedirs(os.path.dirname(out) or ".", exist_ok=True)

        doc = Document()
        normal = doc.styles["Normal"]
        normal.font.name = "Noto Sans SC"
        normal.font.size = Pt(11)

        # Title page.
        h = doc.add_heading(self.title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if self.authors:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(", ".join(self.authors))
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        doc.add_page_break()

        # Flow diagram (embedded PNG).
        doc.add_heading("PRISMA 2020 Flow Diagram", level=1)
        gen = PRISMAFlowGenerator(self.counts, title=self.title, extension="standard")
        with tempfile.NamedTemporaryFile(
            mode="w+b", suffix=".png", delete=False,
        ) as tmp_fig:
            flow_png = tmp_fig.name
        try:
            gen.render_png(flow_png, dpi=150)
            if os.path.exists(flow_png):
                doc.add_picture(flow_png, width=Inches(6.0))
        finally:
            try:
                os.unlink(flow_png)
            except OSError:
                pass
        doc.add_page_break()

        # Checklist.
        doc.add_heading("PRISMA 2020 Checklist", level=1)
        doc.add_paragraph(
            f"Completion: {self.checklist.completion_rate()*100:.1f}% "
            f"({sum(1 for i in self.checklist.items if i.reported)}/"
            f"{len(self.checklist.items)} items reported)."
        )
        ck_table = doc.add_table(rows=1, cols=5)
        ck_table.style = "Light Grid Accent 1"
        hdr = ck_table.rows[0].cells
        for i, label in enumerate(["ID", "Section", "Item", "Location", "Reported"]):
            hdr[i].text = label
            for para in hdr[i].paragraphs:
                for r in para.runs:
                    r.bold = True
        for it in self.checklist.items:
            row = ck_table.add_row().cells
            row[0].text = str(it.id)
            row[1].text = it.section
            row[2].text = it.item_text
            row[3].text = it.location_in_report or "—"
            row[4].text = "✓" if it.reported else "—"
        doc.add_page_break()

        # Included studies.
        doc.add_heading("Included Studies", level=1)
        if not self.extraction:
            doc.add_paragraph("No extraction forms provided.")
        else:
            ex_table = doc.add_table(rows=1, cols=6)
            ex_table.style = "Light Grid Accent 1"
            hdr = ex_table.rows[0].cells
            for i, label in enumerate(
                ["ID", "Design", "Population", "Intervention", "Comparator", "N"],
            ):
                hdr[i].text = label
                for para in hdr[i].paragraphs:
                    for r in para.runs:
                        r.bold = True
            for f in self.extraction:
                row = ex_table.add_row().cells
                row[0].text = f.study_id or "—"
                row[1].text = f.study_design or "—"
                row[2].text = f.population or "—"
                row[3].text = f.intervention or "—"
                row[4].text = f.comparator or "—"
                n = str(f.sample_size) if f.sample_size is not None else "—"
                row[5].text = n

        # Summary.
        doc.add_heading("Summary", level=1)
        c = self.counts
        summary_items = [
            ("Records identified from databases", c.n_records_databases),
            ("Records identified from registers", c.n_records_registers),
            ("Duplicates removed",                c.n_duplicates_removed),
            ("Records screened",                  c.n_records_screened),
            ("Records excluded at title/abstract", c.n_records_excluded_title_abstract),
            ("Full-text assessed",                c.n_full_text_assessed),
            ("Studies in qualitative synthesis",  c.n_studies_included_qualitative),
            ("Studies in quantitative synthesis", c.n_studies_included_quantitative),
        ]
        for label, val in summary_items:
            doc.add_paragraph(f"{label}: {val if val is not None else '—'}")
        doc.add_paragraph(
            f"Checklist completion: {self.checklist.completion_rate()*100:.1f}%"
        )
        doc.add_paragraph(
            f"Included studies: {len(self.extraction)}"
        )

        doc.save(out)
        logger.info("PRISMA report DOCX written to %s", out)
        return out

    # ------------------------------------------------------------------
    # HTML generation.
    # ------------------------------------------------------------------
    def _generate_html(self, output_path: str) -> str:
        """Render the full report as a standalone HTML page.

        Embeds the PRISMA flow diagram as an inline SVG with hover
        tooltips (delegating to :meth:`PRISMAFlowGenerator.render_html`
        for the diagram itself), and a checklist / extraction table in
        HTML.
        """
        out = os.path.abspath(output_path)
        os.makedirs(os.path.dirname(out) or ".", exist_ok=True)

        # Delegate the flow-diagram HTML to the generator.
        with tempfile.NamedTemporaryFile(
            mode="w+", suffix=".html", delete=False, encoding="utf-8",
        ) as tmp_html:
            flow_html_path = tmp_html.name
        try:
            gen = PRISMAFlowGenerator(self.counts, title=self.title, extension="standard")
            gen.render_html(flow_html_path)
            with open(flow_html_path, "r", encoding="utf-8") as f:
                flow_html = f.read()
            # Extract just the inline SVG block (between <div class="prisma-svg"> and </div>).
            svg_start = flow_html.find('<div class="prisma-svg">')
            svg_end   = flow_html.find('</div>', svg_start)
            svg_block = (
                flow_html[svg_start:svg_end + len('</div>')]
                if svg_start >= 0 and svg_end > svg_start
                else "<p>(SVG unavailable)</p>"
            )
        finally:
            try:
                os.unlink(flow_html_path)
            except OSError:
                pass

        # Build the checklist HTML table.
        ck_rows = []
        for it in self.checklist.items:
            rep = "✓" if it.reported else "—"
            loc = _html_escape(it.location_in_report) or "—"
            ck_rows.append(
                f"<tr><td>{it.id}</td><td>{_html_escape(it.section)}</td>"
                f"<td>{_html_escape(it.item_text)}</td>"
                f"<td>{loc}</td><td style='text-align:center'>{rep}</td></tr>"
            )

        # Build the extraction table HTML.
        ex_rows = []
        for f in self.extraction:
            n = f.sample_size if f.sample_size is not None else "—"
            ex_rows.append(
                f"<tr><td>{_html_escape(f.study_id)}</td>"
                f"<td>{_html_escape(f.study_design)}</td>"
                f"<td>{_html_escape(f.population)}</td>"
                f"<td>{_html_escape(f.intervention)}</td>"
                f"<td>{_html_escape(f.comparator)}</td>"
                f"<td style='text-align:right'>{n}</td></tr>"
            )

        authors_html = (
            "<p class='authors'>By " + ", ".join(_html_escape(a) for a in self.authors) + "</p>"
            if self.authors else ""
        )

        completion = self.checklist.completion_rate() * 100
        n_reported = sum(1 for i in self.checklist.items if i.reported)

        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>{_html_escape(self.title)}</title>
<style>
  body {{ font-family: 'Noto Sans SC','DejaVu Sans',Helvetica,Arial,sans-serif;
          margin: 0; padding: 32px 48px; color: #222; line-height: 1.55; }}
  h1 {{ color: #1a3a5c; font-size: 28px; margin-top: 0; }}
  h2 {{ color: #1a3a5c; font-size: 20px; border-bottom: 2px solid #1a3a5c;
        padding-bottom: 4px; margin-top: 32px; }}
  .authors {{ color: #555; font-size: 14px; font-style: italic; }}
  .meta {{ color: #888; font-size: 12px; }}
  table {{ border-collapse: collapse; width: 100%; margin: 12px 0; font-size: 13px; }}
  th, td {{ border: 1px solid #ccc; padding: 6px 8px; text-align: left;
            vertical-align: top; }}
  th {{ background: #1a3a5c; color: #fff; }}
  tr:nth-child(even) td {{ background: #f4f7fb; }}
  .prisma-svg {{ margin: 16px 0; max-width: 900px; }}
  .prisma-svg svg {{ width: 100%; height: auto; }}
  .summary {{ background: #f4f7fb; border-left: 4px solid #1a3a5c;
              padding: 12px 16px; margin: 12px 0; }}
  .prisma-svg svg rect.flow-box,
  .prisma-svg svg rect.exclusion-box {{ cursor: help; }}
  .prisma-svg svg rect.flow-box:hover,
  .prisma-svg svg rect.exclusion-box:hover {{
     stroke-width: 3; filter: drop-shadow(0 0 6px rgba(46,92,138,0.5)); }}
  footer {{ margin-top: 40px; padding-top: 12px; border-top: 1px solid #ccc;
            color: #666; font-size: 11px; text-align: center; font-style: italic; }}
</style>
</head>
<body>
  <h1>{_html_escape(self.title)}</h1>
  {authors_html}
  <p class="meta">PRISMA 2020 reporting package —
     Page MJ et al., <i>BMJ</i> 2021;372:n71.</p>

  <h2>PRISMA 2020 Flow Diagram</h2>
  {svg_block}

  <h2>PRISMA 2020 Checklist</h2>
  <div class="summary">
    <strong>Completion:</strong> {completion:.1f}%
    ({n_reported}/{len(self.checklist.items)} items reported)
  </div>
  <table>
    <thead><tr>
      <th style="width:5%">ID</th>
      <th style="width:12%">Section</th>
      <th style="width:55%">Item</th>
      <th style="width:23%">Location in report</th>
      <th style="width:5%">Reported</th>
    </tr></thead>
    <tbody>
      {''.join(ck_rows)}
    </tbody>
  </table>

  <h2>Included Studies</h2>
  <p>{len(self.extraction)} study extraction form(s) included.</p>
  <table>
    <thead><tr>
      <th>Study ID</th><th>Design</th><th>Population</th>
      <th>Intervention</th><th>Comparator</th><th style="text-align:right">N</th>
    </tr></thead>
    <tbody>
      {''.join(ex_rows) if ex_rows else '<tr><td colspan="6"><em>No extraction forms provided.</em></td></tr>'}
    </tbody>
  </table>

  <h2>Summary</h2>
  <div class="summary">
    {_html_escape(self._summary_text())}
  </div>

  <footer>Generated by Academic Research Suite v2.0.0 — PRISMA 2020
  (Page MJ et al., BMJ 2021;372:n71).</footer>
</body>
</html>
"""
        with open(out, "w", encoding="utf-8") as f:
            f.write(html)
        logger.info("PRISMA report HTML written to %s", out)
        return out

    # ------------------------------------------------------------------
    # Helpers.
    # ------------------------------------------------------------------
    def _summary_text(self) -> str:
        """Plain-text summary for embedding in HTML/PDF body."""
        c = self.counts
        lines = [
            f"Records identified from databases: {c.n_records_databases if c.n_records_databases is not None else '—'}",
            f"Records identified from registers: {c.n_records_registers if c.n_records_registers is not None else '—'}",
            f"Duplicates removed: {c.n_duplicates_removed if c.n_duplicates_removed is not None else '—'}",
            f"Records screened: {c.n_records_screened if c.n_records_screened is not None else '—'}",
            f"Records excluded at title/abstract: {c.n_records_excluded_title_abstract if c.n_records_excluded_title_abstract is not None else '—'}",
            f"Full-text assessed: {c.n_full_text_assessed if c.n_full_text_assessed is not None else '—'}",
            f"Studies in qualitative synthesis: {c.n_studies_included_qualitative if c.n_studies_included_qualitative is not None else '—'}",
            f"Studies in quantitative synthesis: {c.n_studies_included_quantitative if c.n_studies_included_quantitative is not None else '—'}",
            f"Checklist completion: {self.checklist.completion_rate()*100:.1f}%",
            f"Included studies: {len(self.extraction)}",
        ]
        return "<br>".join(lines)


# ---------------------------------------------------------------------------
# Module-level helpers.
# ---------------------------------------------------------------------------
def _html_escape(text: str) -> str:
    """Escape HTML-significant characters."""
    if not text:
        return ""
    return (str(text)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;"))


__all__ = ["PRISMAReport"]
