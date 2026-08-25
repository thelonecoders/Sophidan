"""Study-protocol templates and a protocol-builder helper.

This module ships pre-built study-protocol templates aligned with major
reporting guidelines:

* :meth:`ProtocolTemplateLibrary.systematic_review` — PRISMA 2020-aligned
  SR protocol (14 sections).
* :meth:`ProtocolTemplateLibrary.scoping_review` — JBI / Arksey & O'Malley
  scoping-review protocol.
* :meth:`ProtocolTemplateLibrary.meta_analysis_protocol` — Cochrane-style
  meta-analysis protocol.
* :meth:`ProtocolTemplateLibrary.rapid_review` — WHO / Cochrane Rapid Review
  Methods.
* :meth:`ProtocolTemplateLibrary.case_study_protocol` — Yin-style case-study
  protocol.
* :meth:`ProtocolTemplateLibrary.cohort_study_protocol` — STROBE-aligned
  observational cohort protocol.
* :meth:`ProtocolTemplateLibrary.rct_protocol` — CONSORT / SPIRIT-aligned
  RCT protocol.
* :meth:`ProtocolTemplateLibrary.qualitative_protocol` — ENTREQ-aligned
  qualitative-study protocol.
* :meth:`ProtocolTemplateLibrary.mixed_methods_protocol` — MMAT-aligned
  mixed-methods protocol.

A :class:`ProtocolBuilder` turns a :class:`ProtocolTemplate` into a
:class:`Protocol`, validates that all required sections are filled,
renders the protocol to Markdown / PDF / DOCX, and supports per-section
filling with arbitrary body text.

Heavy deps (``reportlab``, ``python-docx``) are imported lazily inside the
export methods so the module is importable in minimal environments.
"""
#
# MIT License — Academic Research Suite
# Copyright (c) 2026 — see /LICENSE for full text.
#
from __future__ import annotations

import logging
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import List, Optional, Tuple

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------
@dataclass
class ProtocolSection:
    """A single section in a study-protocol template.

    Attributes:
        heading: Section heading (e.g. ``"Background"``).
        body: Default body text (often a placeholder).
        required: Whether this section must be filled before the protocol
            can be considered complete (used by
            :meth:`ProtocolBuilder.validate`).
        placeholder_text: Hint shown to the user when the section is
            unfilled.
        hints: A list of writing hints / prompts.
    """

    heading: str
    body: str = ""
    required: bool = False
    placeholder_text: str = ""
    hints: List[str] = field(default_factory=list)


@dataclass
class ProtocolTemplate:
    """A reusable study-protocol template.

    Attributes:
        name: Short identifier (e.g. ``"systematic_review"``).
        sections: Ordered list of :class:`ProtocolSection`.
    """

    name: str
    sections: List[ProtocolSection] = field(default_factory=list)


@dataclass
class Protocol:
    """An instantiated, fillable protocol.

    Attributes:
        template_name: Name of the source :class:`ProtocolTemplate`.
        sections: Ordered ``(heading, body)`` tuples.
        created_at: ISO-8601 UTC timestamp.
        version: Protocol version string (semver-ish).
        project_id: Optional numeric project identifier.
    """

    template_name: str
    sections: List[Tuple[str, str]] = field(default_factory=list)
    created_at: str = field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )
    version: str = "1.0.0"
    project_id: Optional[int] = None

    def section(self, heading: str) -> Optional[Tuple[str, str]]:
        """Return the ``(heading, body)`` tuple for ``heading``, or None."""
        for h, b in self.sections:
            if h.lower() == heading.lower():
                return (h, b)
        return None


# ---------------------------------------------------------------------------
# Template library
# ---------------------------------------------------------------------------
class ProtocolTemplateLibrary:
    """Factory of pre-built study-protocol templates.

    Each classmethod returns a fully populated :class:`ProtocolTemplate`
    ready to be handed to :class:`ProtocolBuilder.from_template`.
    """

    # ------------------------------------------------------------------
    # Templates
    # ------------------------------------------------------------------
    @classmethod
    def systematic_review(cls) -> ProtocolTemplate:
        """Return a PRISMA 2020-aligned systematic-review protocol (14 sections)."""
        sections = [
            ProtocolSection(
                "Title",
                required=True,
                placeholder_text="<Descriptive title of the review>",
                hints=[
                    "Include PICO(S) elements where applicable.",
                    "End with ': a systematic review'.",
                ],
            ),
            ProtocolSection(
                "Authors and affiliations",
                required=True,
                placeholder_text="<Names, ORCID, affiliations>",
                hints=["List all contributing authors."],
            ),
            ProtocolSection(
                "Background and rationale",
                required=True,
                placeholder_text="<Why this review is needed>",
                hints=[
                    "Frame the problem and the gap in prior reviews.",
                    "Cite prior related reviews and primary studies.",
                ],
            ),
            ProtocolSection(
                "Objectives",
                required=True,
                placeholder_text="<Primary and secondary objectives>",
                hints=["State objectives as PICO questions."],
            ),
            ProtocolSection(
                "Eligibility criteria",
                required=True,
                placeholder_text="<Inclusion / exclusion criteria>",
                hints=[
                    "Specify population, intervention, comparator, "
                    "outcomes, study designs.",
                    "List language / date / publication-status "
                    "restrictions.",
                ],
            ),
            ProtocolSection(
                "Information sources",
                required=True,
                placeholder_text="<Databases, registers, hand-searches>",
                hints=[
                    "List databases (≥3 recommended) + last search date.",
                    "List trial registers, grey literature sources.",
                ],
            ),
            ProtocolSection(
                "Search strategy",
                required=True,
                placeholder_text="<Full Boolean search strings>",
                hints=[
                    "Provide at least one full reproducible search "
                    "string per database.",
                ],
            ),
            ProtocolSection(
                "Selection process",
                required=True,
                placeholder_text="<Screening workflow>",
                hints=[
                    "Number of independent screeners (≥2 recommended).",
                    "Conflict-resolution mechanism.",
                    "Tool used (e.g. Covidence, Rayyan).",
                ],
            ),
            ProtocolSection(
                "Data collection process",
                required=True,
                placeholder_text="<Extraction form & piloting>",
                hints=[
                    "Pilot the form on ≥5 studies.",
                    "State whether extractors are blinded.",
                ],
            ),
            ProtocolSection(
                "Data items",
                required=True,
                placeholder_text="<List of variables>",
                hints=[
                    "List all extracted variables, prioritise PICO.",
                    "Note any assumptions / simplifications.",
                ],
            ),
            ProtocolSection(
                "Outcomes and prioritisation",
                required=True,
                placeholder_text="<Primary, secondary, exploratory>",
                hints=[
                    "Distinguish primary vs secondary outcomes.",
                ],
            ),
            ProtocolSection(
                "Risk of bias assessment",
                required=True,
                placeholder_text="<RoB tool per study design>",
                hints=[
                    "Pick RoB 2 (RCTs), ROBINS-I (non-randomised), "
                    "QUADAS-2 (diagnostic accuracy).",
                ],
            ),
            ProtocolSection(
                "Data synthesis",
                required=True,
                placeholder_text="<Narrative and/or meta-analysis plan>",
                hints=[
                    "State whether meta-analysis is planned and the "
                    "effect measure (RR, OR, SMD, etc.).",
                    "Heterogeneity assessment (I², τ², Q).",
                    "Subgroup & sensitivity analyses.",
                ],
            ),
            ProtocolSection(
                "Certainty assessment",
                required=False,
                placeholder_text="<GRADE or alternative>",
                hints=["Use GRADE for body-of-evidence certainty."],
            ),
            ProtocolSection(
                "Registration and dissemination",
                required=False,
                placeholder_text="<PROSPERO ID; publication plan>",
                hints=[
                    "Register on PROSPERO before screening.",
                    "State dissemination channels.",
                ],
            ),
        ]
        return ProtocolTemplate(name="systematic_review", sections=sections)

    @classmethod
    def scoping_review(cls) -> ProtocolTemplate:
        """Return a JBI / Arksey & O'Malley scoping-review protocol."""
        sections = [
            ProtocolSection(
                "Title",
                required=True,
                placeholder_text="<Title ending with ': a scoping review'>",
            ),
            ProtocolSection(
                "Authors",
                required=True,
                placeholder_text="<Authors / affiliations>",
            ),
            ProtocolSection(
                "Background",
                required=True,
                placeholder_text="<Rationale for choosing a scoping review>",
                hints=[
                    "Cite Arksey & O'Malley (2005), Levac et al. (2010), "
                    "JBI scoping review methodology.",
                ],
            ),
            ProtocolSection(
                "Objectives / review questions",
                required=True,
                placeholder_text="<Broad research questions>",
                hints=[
                    "Scoping reviews answer broad mapping questions; "
                    "avoid narrow PICO.",
                ],
            ),
            ProtocolSection(
                "Inclusion criteria (PCC)",
                required=True,
                placeholder_text="<Population, Concept, Context>",
                hints=[
                    "Use PCC (Population, Concept, Context) rather than "
                    "PICO.",
                ],
            ),
            ProtocolSection(
                "Search strategy",
                required=True,
                placeholder_text="<Databases + sample search string>",
            ),
            ProtocolSection(
                "Evidence selection",
                required=True,
                placeholder_text="<Screening workflow>",
                hints=["≥2 independent reviewers."],
            ),
            ProtocolSection(
                "Data extraction",
                required=True,
                placeholder_text="<Charting form fields>",
            ),
            ProtocolSection(
                "Data analysis and presentation",
                required=True,
                placeholder_text="<Charting, numerical, thematic>",
                hints=[
                    "Arksey & O'Malley framework + optional stakeholder "
                    "consultation.",
                ],
            ),
            ProtocolSection(
                "Stakeholder consultation (optional)",
                required=False,
                placeholder_text="<Patient / public involvement>",
            ),
            ProtocolSection(
                "Reporting",
                required=False,
                placeholder_text="<PRISMA-ScR adherence>",
                hints=["Use PRISMA-ScR checklist for reporting."],
            ),
        ]
        return ProtocolTemplate(name="scoping_review", sections=sections)

    @classmethod
    def meta_analysis_protocol(cls) -> ProtocolTemplate:
        """Return a Cochrane-style meta-analysis protocol."""
        sections = [
            ProtocolSection("Title", required=True, placeholder_text="<Title>"),
            ProtocolSection("Authors", required=True, placeholder_text="<Authors>"),
            ProtocolSection(
                "Background",
                required=True,
                placeholder_text="<Rationale + prior reviews>",
            ),
            ProtocolSection(
                "Objectives",
                required=True,
                placeholder_text="<Primary and secondary objectives>",
            ),
            ProtocolSection(
                "Criteria for considering studies",
                required=True,
                placeholder_text="<Types of studies, participants, "
                                  "interventions, outcomes>",
            ),
            ProtocolSection(
                "Search methods",
                required=True,
                placeholder_text="<Databases, search strings>",
            ),
            ProtocolSection(
                "Data collection and analysis",
                required=True,
                placeholder_text="<Selection, extraction, RoB>",
            ),
            ProtocolSection(
                "Effect measures",
                required=True,
                placeholder_text="<RR, OR, SMD, hazard ratio, etc.>",
            ),
            ProtocolSection(
                "Unit of analysis issues",
                required=True,
                placeholder_text="<Cluster, crossover, multi-arm>",
            ),
            ProtocolSection(
                "Dealing with missing data",
                required=True,
                placeholder_text="<Imputation strategy>",
            ),
            ProtocolSection(
                "Assessment of heterogeneity",
                required=True,
                placeholder_text="<I², τ², visual inspection>",
            ),
            ProtocolSection(
                "Assessment of reporting biases",
                required=True,
                placeholder_text="<Funnel plot, Egger test>",
            ),
            ProtocolSection(
                "Data synthesis",
                required=True,
                placeholder_text="<Random-effects model, software>",
                hints=["Use DerSimonian-Laird or REML estimator."],
            ),
            ProtocolSection(
                "Subgroup and sensitivity analyses",
                required=True,
                placeholder_text="<Pre-specified analyses>",
            ),
            ProtocolSection(
                "Certainty of evidence",
                required=True,
                placeholder_text="<GRADE>",
            ),
            ProtocolSection(
                "Acknowledgements / funding",
                required=False,
                placeholder_text="<Funders, contributors>",
            ),
        ]
        return ProtocolTemplate(
            name="meta_analysis_protocol", sections=sections
        )

    @classmethod
    def rapid_review(cls) -> ProtocolTemplate:
        """Return a WHO / Cochrane Rapid Review protocol."""
        sections = [
            ProtocolSection("Title", required=True, placeholder_text="<Title>"),
            ProtocolSection(
                "Background",
                required=True,
                placeholder_text="<Urgency rationale>",
                hints=["Explain why a rapid review is justified."],
            ),
            ProtocolSection("Objectives", required=True, placeholder_text="<Objectives>"),
            ProtocolSection(
                "Eligibility criteria",
                required=True,
                placeholder_text="<Streamlined PICO>",
            ),
            ProtocolSection(
                "Search strategy",
                required=True,
                placeholder_text="<2-3 databases; date-limited>",
                hints=[
                    "Restrict to 2-3 databases and to recent years "
                    "where appropriate.",
                ],
            ),
            ProtocolSection(
                "Screening",
                required=True,
                placeholder_text="<Single-reviewer + spot-check>",
                hints=[
                    "Single-reviewer screening with second-reviewer "
                    "verification of exclusions is acceptable.",
                ],
            ),
            ProtocolSection(
                "Data extraction",
                required=True,
                placeholder_text="<Limited variables>",
            ),
            ProtocolSection(
                "Risk of bias",
                required=True,
                placeholder_text="<Streamlined RoB>",
            ),
            ProtocolSection(
                "Synthesis",
                required=True,
                placeholder_text="<Narrative or limited meta-analysis>",
            ),
            ProtocolSection(
                "Reporting & limitations",
                required=True,
                placeholder_text="<Cochrane Rapid Review methods>",
            ),
        ]
        return ProtocolTemplate(name="rapid_review", sections=sections)

    @classmethod
    def case_study_protocol(cls) -> ProtocolTemplate:
        """Return a Yin-style case-study protocol."""
        sections = [
            ProtocolSection(
                "Title", required=True, placeholder_text="<Case title>"
            ),
            ProtocolSection(
                "Background",
                required=True,
                placeholder_text="<Phenomenon and case context>",
            ),
            ProtocolSection(
                "Research questions / propositions",
                required=True,
                placeholder_text="<Propositions to test>",
                hints=["State explicit propositions (Yin 2018)."],
            ),
            ProtocolSection(
                "Unit of analysis",
                required=True,
                placeholder_text="<Individual / group / organisation>",
            ),
            ProtocolSection(
                "Case selection criteria",
                required=True,
                placeholder_text="<Why this case>"),
            ProtocolSection(
                "Data sources (triangulation)",
                required=True,
                placeholder_text="<Documents, interviews, observations>",
            ),
            ProtocolSection(
                "Data collection procedures",
                required=True,
                placeholder_text="<Timeline, instruments>",
            ),
            ProtocolSection(
                "Data analysis",
                required=True,
                placeholder_text="<Pattern matching, explanation building>",
            ),
            ProtocolSection(
                "Validity / trustworthiness",
                required=True,
                placeholder_text="<Construct, internal, external validity>",
            ),
            ProtocolSection(
                "Ethical considerations",
                required=True,
                placeholder_text="<IRB, consent, anonymisation>",
            ),
            ProtocolSection(
                "Reporting",
                required=False,
                placeholder_text="<CARE or equivalent>",
            ),
        ]
        return ProtocolTemplate(name="case_study_protocol", sections=sections)

    @classmethod
    def cohort_study_protocol(cls) -> ProtocolTemplate:
        """Return a STROBE-aligned cohort-study protocol."""
        sections = [
            ProtocolSection("Title", required=True, placeholder_text="<Title>"),
            ProtocolSection(
                "Background and rationale",
                required=True,
                placeholder_text="<Background>",
            ),
            ProtocolSection(
                "Objectives and hypotheses",
                required=True,
                placeholder_text="<Primary / secondary hypotheses>",
            ),
            ProtocolSection(
                "Study design",
                required=True,
                placeholder_text="<Prospective / retrospective cohort>",
            ),
            ProtocolSection(
                "Setting",
                required=True,
                placeholder_text="<Location, period>",
            ),
            ProtocolSection(
                "Participants",
                required=True,
                placeholder_text="<Eligibility, recruitment>",
            ),
            ProtocolSection(
                "Variables",
                required=True,
                placeholder_text="<Exposure, outcome, confounders, mediators>",
            ),
            ProtocolSection(
                "Data sources / measurements",
                required=True,
                placeholder_text="<Measurement methods>",
            ),
            ProtocolSection(
                "Bias",
                required=True,
                placeholder_text="<Selection, information, confounding>",
            ),
            ProtocolSection(
                "Study size",
                required=True,
                placeholder_text="<Sample size calculation>",
            ),
            ProtocolSection(
                "Quantitative variables",
                required=True,
                placeholder_text="<Continuous / categorical handling>",
            ),
            ProtocolSection(
                "Statistical methods",
                required=True,
                placeholder_text="<Regression, survival analysis, "
                                  "missing data>",
            ),
            ProtocolSection(
                "Ethics",
                required=True,
                placeholder_text="<IRB approval, consent>",
            ),
            ProtocolSection(
                "Dissemination",
                required=False,
                placeholder_text="<Publication plan>",
            ),
        ]
        return ProtocolTemplate(name="cohort_study_protocol", sections=sections)

    @classmethod
    def rct_protocol(cls) -> ProtocolTemplate:
        """Return a CONSORT + SPIRIT-aligned RCT protocol."""
        sections = [
            ProtocolSection(
                "Title",
                required=True,
                placeholder_text="<Title with design descriptor>",
                hints=["Use SPIRIT 2013 title format."],
            ),
            ProtocolSection(
                "Trial registration",
                required=True,
                placeholder_text="<ClinicalTrials.gov / ISRCTN ID>",
            ),
            ProtocolSection(
                "Rationale and background",
                required=True,
                placeholder_text="<Background>",
            ),
            ProtocolSection(
                "Objectives and hypotheses",
                required=True,
                placeholder_text="<Primary / secondary>",
            ),
            ProtocolSection(
                "Trial design",
                required=True,
                placeholder_text="<Parallel / cluster / crossover>",
            ),
            ProtocolSection(
                "Study setting",
                required=True,
                placeholder_text="<Sites, country>",
            ),
            ProtocolSection(
                "Eligibility criteria",
                required=True,
                placeholder_text="<Inclusion / exclusion>",
            ),
            ProtocolSection(
                "Interventions",
                required=True,
                placeholder_text="<Description, comparator>",
            ),
            ProtocolSection(
                "Outcomes",
                required=True,
                placeholder_text="<Primary, secondary>",
            ),
            ProtocolSection(
                "Participant timeline",
                required=True,
                placeholder_text="<Schedule of enrolment, interventions, "
                                  "assessments>",
            ),
            ProtocolSection(
                "Sample size",
                required=True,
                placeholder_text="<Calculation + assumptions>",
            ),
            ProtocolSection(
                "Recruitment",
                required=True,
                placeholder_text="<Strategies>",
            ),
            ProtocolSection(
                "Allocation",
                required=True,
                placeholder_text="<Sequence generation, concealment, "
                                  "implementation>",
            ),
            ProtocolSection(
                "Blinding",
                required=True,
                placeholder_text="<Who is blinded>",
            ),
            ProtocolSection(
                "Data collection methods",
                required=True,
                placeholder_text="<Instruments, retention>",
            ),
            ProtocolSection(
                "Statistical methods",
                required=True,
                placeholder_text="<Primary, secondary, interim>",
            ),
            ProtocolSection(
                "Data monitoring",
                required=True,
                placeholder_text="<DMC, harms, stopping rules>",
            ),
            ProtocolSection(
                "Harms",
                required=True,
                placeholder_text="<Adverse event collection>",
            ),
            ProtocolSection(
                "Auditing",
                required=False,
                placeholder_text="<Audit plan>",
            ),
            ProtocolSection(
                "Research ethics approval",
                required=True,
                placeholder_text="<IRB>",
            ),
            ProtocolSection(
                "Protocol amendments",
                required=False,
                placeholder_text="<Amendment procedure>",
            ),
            ProtocolSection(
                "Confidentiality / ancillary / post-trial care",
                required=False,
                placeholder_text="<Provisions>",
            ),
            ProtocolSection(
                "Dissemination policy",
                required=False,
                placeholder_text="<Publication, data sharing>",
            ),
            ProtocolSection(
                "Informed consent",
                required=True,
                placeholder_text="<Consent process>",
            ),
            ProtocolSection(
                "Funding / sponsorship / conflicts",
                required=False,
                placeholder_text="<Funding + COI>",
            ),
        ]
        return ProtocolTemplate(name="rct_protocol", sections=sections)

    @classmethod
    def qualitative_protocol(cls) -> ProtocolTemplate:
        """Return an ENTREQ-aligned qualitative-study protocol."""
        sections = [
            ProtocolSection(
                "Title", required=True, placeholder_text="<Title>"
            ),
            ProtocolSection(
                "Authors / affiliations",
                required=True,
                placeholder_text="<Authors>",
            ),
            ProtocolSection(
                "Background and rationale",
                required=True,
                placeholder_text="<Phenomenon of interest>",
            ),
            ProtocolSection(
                "Research question(s)",
                required=True,
                placeholder_text="<Broad qualitative question>",
            ),
            ProtocolSection(
                "Methodology",
                required=True,
                placeholder_text="<Phenomenology / grounded theory / "
                                  "ethnography / framework>",
            ),
            ProtocolSection(
                "Sampling strategy",
                required=True,
                placeholder_text="<Purposive / theoretical / snowball>",
            ),
            ProtocolSection(
                "Participants and setting",
                required=True,
                placeholder_text="<Who / where>",
            ),
            ProtocolSection(
                "Data collection",
                required=True,
                placeholder_text="<Interviews / focus groups / observations>",
                hints=["State the interviewer-interviewee relationship."],
            ),
            ProtocolSection(
                "Interview guide",
                required=True,
                placeholder_text="<Topics / questions>",
            ),
            ProtocolSection(
                "Data analysis",
                required=True,
                placeholder_text="<Thematic / framework / grounded>",
                hints=[
                    "State software (NVivo, ATLAS.ti) and saturation "
                    "criteria.",
                ],
            ),
            ProtocolSection(
                "Trustworthiness",
                required=True,
                placeholder_text="<Credibility, transferability, "
                                  "dependability, confirmability>",
            ),
            ProtocolSection(
                "Ethical considerations",
                required=True,
                placeholder_text="<IRB, consent>",
            ),
            ProtocolSection(
                "Reflexivity",
                required=True,
                placeholder_text="<Researcher positionality>",
            ),
            ProtocolSection(
                "Dissemination",
                required=False,
                placeholder_text="<Channels>",
            ),
        ]
        return ProtocolTemplate(name="qualitative_protocol", sections=sections)

    @classmethod
    def mixed_methods_protocol(cls) -> ProtocolTemplate:
        """Return an MMAT-aligned mixed-methods protocol."""
        sections = [
            ProtocolSection("Title", required=True, placeholder_text="<Title>"),
            ProtocolSection(
                "Background",
                required=True,
                placeholder_text="<Rationale for mixing methods>",
            ),
            ProtocolSection(
                "Research questions",
                required=True,
                placeholder_text="<Quantitative + qualitative questions>",
            ),
            ProtocolSection(
                "Design",
                required=True,
                placeholder_text="<Convergent / explanatory / exploratory>",
                hints=[
                    "Reference Creswell & Plano Clark (2018).",
                ],
            ),
            ProtocolSection(
                "Quantitative component",
                required=True,
                placeholder_text="<Design, sample, measures>",
            ),
            ProtocolSection(
                "Qualitative component",
                required=True,
                placeholder_text="<Design, sample, data>",
            ),
            ProtocolSection(
                "Integration",
                required=True,
                placeholder_text="<Point of integration: design, methods, "
                                  "interpretation>",
            ),
            ProtocolSection(
                "Mixed-methods appraisal (MMAT)",
                required=True,
                placeholder_text="<Quality criteria per MMAT>",
            ),
            ProtocolSection(
                "Sampling and recruitment",
                required=True,
                placeholder_text="<Joint display plans>",
            ),
            ProtocolSection(
                "Data analysis",
                required=True,
                placeholder_text="<Quant analysis, qual analysis, joint "
                                  "display>",
            ),
            ProtocolSection(
                "Rigor / validity",
                required=True,
                placeholder_text="<Legitimation criteria>",
            ),
            ProtocolSection(
                "Ethics",
                required=True,
                placeholder_text="<IRB, consent>",
            ),
            ProtocolSection(
                "Dissemination",
                required=False,
                placeholder_text="<Plans>",
            ),
        ]
        return ProtocolTemplate(name="mixed_methods_protocol", sections=sections)

    # ------------------------------------------------------------------
    # Library accessor API
    # ------------------------------------------------------------------
    @classmethod
    def available(cls) -> List[str]:
        """Return the names of every available template."""
        return [
            "systematic_review",
            "scoping_review",
            "meta_analysis_protocol",
            "rapid_review",
            "case_study_protocol",
            "cohort_study_protocol",
            "rct_protocol",
            "qualitative_protocol",
            "mixed_methods_protocol",
        ]

    @classmethod
    def get(cls, name: str) -> ProtocolTemplate:
        """Look up a template by name (case-insensitive).

        Args:
            name: Template name (e.g. ``"systematic_review"``).

        Returns:
            A :class:`ProtocolTemplate`.

        Raises:
            KeyError: If the name is not in :meth:`available`.
        """
        key = (name or "").strip().lower()
        mapping = {
            "systematic_review": cls.systematic_review,
            "scoping_review": cls.scoping_review,
            "meta_analysis_protocol": cls.meta_analysis_protocol,
            "rapid_review": cls.rapid_review,
            "case_study_protocol": cls.case_study_protocol,
            "cohort_study_protocol": cls.cohort_study_protocol,
            "rct_protocol": cls.rct_protocol,
            "qualitative_protocol": cls.qualitative_protocol,
            "mixed_methods_protocol": cls.mixed_methods_protocol,
        }
        if key not in mapping:
            raise KeyError(
                f"Unknown protocol template: {name!r}. "
                f"Available: {cls.available()}"
            )
        return mapping[key]()


# ---------------------------------------------------------------------------
# Protocol builder
# ---------------------------------------------------------------------------
class ProtocolBuilder:
    """Build, fill, validate and export :class:`Protocol` instances."""

    @staticmethod
    def from_template(
        template: ProtocolTemplate, project_id: Optional[int] = None
    ) -> Protocol:
        """Instantiate a :class:`Protocol` from a :class:`ProtocolTemplate`.

        The protocol's sections are pre-populated with the template's
        default body text (which is often a placeholder — callers replace
        it via :meth:`fill_section`).
        """
        return Protocol(
            template_name=template.name,
            sections=[(s.heading, s.body) for s in template.sections],
            project_id=project_id,
        )

    @staticmethod
    def fill_section(
        protocol: Protocol, heading: str, body: str
    ) -> Protocol:
        """Replace the body of the section with ``heading``.

        If the heading is not already present in ``protocol.sections``,
        the section is appended.

        Returns:
            The *same* protocol (mutated in-place; returned for chaining).
        """
        body = body or ""
        for i, (h, _) in enumerate(protocol.sections):
            if h.lower() == heading.lower():
                protocol.sections[i] = (h, body)
                return protocol
        protocol.sections.append((heading, body))
        return protocol

    @staticmethod
    def validate(protocol: Protocol) -> List[str]:
        """Return a list of missing-required-section headings.

        A section is "missing" if its body is empty OR equals its
        template placeholder text.  Returns an empty list when the
        protocol is complete.
        """
        template: Optional[ProtocolTemplate] = None
        try:
            template = ProtocolTemplateLibrary.get(protocol.template_name)
        except KeyError:
            return []
        missing: List[str] = []
        # Index sections by lowercased heading for quick lookup.
        sec_map = {h.lower(): b for h, b in protocol.sections}
        for tpl_sec in template.sections:
            if not tpl_sec.required:
                continue
            body = sec_map.get(tpl_sec.heading.lower(), "")
            if not body or body == tpl_sec.placeholder_text:
                missing.append(tpl_sec.heading)
        return missing

    @staticmethod
    def to_markdown(protocol: Protocol) -> str:
        """Render the protocol as a Markdown document string."""
        lines: List[str] = []
        title = protocol.template_name.replace("_", " ").title()
        lines.append(f"# {title} Protocol")
        lines.append("")
        lines.append(
            f"- **Template:** `{protocol.template_name}`"
        )
        lines.append(f"- **Version:** {protocol.version}")
        lines.append(f"- **Created:** {protocol.created_at}")
        if protocol.project_id is not None:
            lines.append(f"- **Project ID:** {protocol.project_id}")
        lines.append("")
        lines.append("---")
        lines.append("")
        for heading, body in protocol.sections:
            lines.append(f"## {heading}")
            lines.append("")
            lines.append(body or "_(not yet filled)_")
            lines.append("")
        return "\n".join(lines).rstrip() + "\n"

    @staticmethod
    def to_pdf(protocol: Protocol, path: str) -> str:
        """Export the protocol to a PDF file (lazy ``reportlab`` import).

        Args:
            protocol: The protocol to export.
            path: Output PDF file path.

        Returns:
            The absolute path of the written file.
        """
        try:
            from reportlab.lib.pagesizes import A4  # type: ignore[import]
            from reportlab.lib.styles import getSampleStyleSheet  # type: ignore[import]
            from reportlab.lib.units import cm  # type: ignore[import]
            from reportlab.platypus import (  # type: ignore[import]
                Paragraph,
                SimpleDocTemplate,
                Spacer,
            )
        except ImportError as exc:  # pragma: no cover - depends on env
            raise ImportError(
                "reportlab is required for ProtocolBuilder.to_pdf; "
                "install with: pip install reportlab"
            ) from exc

        import os

        abs_path = os.path.abspath(path)
        doc = SimpleDocTemplate(
            abs_path,
            pagesize=A4,
            leftMargin=2 * cm,
            rightMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
            title=f"{protocol.template_name} protocol",
        )
        styles = getSampleStyleSheet()
        story = [
            Paragraph(
                protocol.template_name.replace("_", " ").title()
                + " Protocol",
                styles["Title"],
            ),
            Spacer(1, 0.5 * cm),
            Paragraph(
                f"Version {protocol.version} — created {protocol.created_at}",
                styles["Normal"],
            ),
            Spacer(1, 0.5 * cm),
        ]
        for heading, body in protocol.sections:
            story.append(Paragraph(heading, styles["Heading2"]))
            text = body or "_(not yet filled)_"
            # Escape HTML-significant characters for safety.
            text = (
                text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            story.append(Paragraph(text.replace("\n", "<br/>"), styles["Normal"]))
            story.append(Spacer(1, 0.3 * cm))
        doc.build(story)
        return abs_path

    @staticmethod
    def to_docx(protocol: Protocol, path: str) -> str:
        """Export the protocol to a .docx file (lazy ``python-docx`` import).

        Args:
            protocol: The protocol to export.
            path: Output .docx file path.

        Returns:
            The absolute path of the written file.
        """
        try:
            from docx import Document  # type: ignore[import]
        except ImportError as exc:  # pragma: no cover - depends on env
            raise ImportError(
                "python-docx is required for ProtocolBuilder.to_docx; "
                "install with: pip install python-docx"
            ) from exc

        import os

        abs_path = os.path.abspath(path)
        doc = Document()
        doc.add_heading(
            protocol.template_name.replace("_", " ").title() + " Protocol",
            level=0,
        )
        doc.add_paragraph(
            f"Version {protocol.version} — created {protocol.created_at}"
        )
        for heading, body in protocol.sections:
            doc.add_heading(heading, level=2)
            doc.add_paragraph(body or "_(not yet filled)_")
        doc.save(abs_path)
        return abs_path


__all__ = [
    "ProtocolSection",
    "ProtocolTemplate",
    "ProtocolTemplateLibrary",
    "ProtocolBuilder",
    "Protocol",
]
