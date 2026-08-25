"""Narrative & tabular meta-analysis report generator.

Produces PRISMA + Cochrane-style reports containing:

* Characteristics-of-included-studies table.
* Forest plot (Cochrane style).
* Funnel plot with publication-bias diagnostics.
* Leave-one-out sensitivity table.
* Subgroup analysis table (when ``subgroups`` provided).
* Summary-of-findings (SoF) table.
* GRADE assessment summary (when ``grade`` settings are provided).

Output formats: Markdown, HTML, PDF, DOCX.
Heavy deps (matplotlib, pandas, reportlab, python-docx, markdown) are
lazy-imported.
"""
#
# MIT License — Academic Research Suite
# Copyright (c) 2026 — see /LICENSE for full text.
#

from __future__ import annotations

import base64
import logging
import os
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Tuple

from .effect_sizes import EffectSize, EffectSizeType
from .pooling import MetaAnalysisResult, PoolingMethod
from .subgroup import SensitivityAnalysis, SubgroupAnalysis
from .forest_plot import ForestPlot
from .funnel_plot import FunnelPlot

logger = logging.getLogger(__name__)

__all__ = ["MetaAnalysisReport"]


@dataclass
class MetaAnalysisReport:
    """Generator for a complete meta-analysis report.

    Attributes:
        meta_result: The :class:`MetaAnalysisResult` to report on.
        effect_sizes: List of per-study :class:`EffectSize`.
        study_data: Optional list of dicts (one per study) with extra
            metadata (e.g. ``{'study_id': 's1', 'authors': 'Smith et al',
            'year': 2018, 'population': 'adults', 'intervention': 'drug A',
            'comparator': 'placebo', 'outcome': 'mortality', 'design': 'RCT'}``).
        settings: Optional dict with keys:
            ``title``, ``authors``, ``subgroups`` (dict study_id→subgroup),
            ``grade`` (dict of GRADE ratings), ``style`` (forest/funnel style).
    """

    meta_result: MetaAnalysisResult
    effect_sizes: List[EffectSize]
    study_data: Optional[List[dict]] = None
    settings: dict = field(default_factory=dict)

    # ------------------------------------------------------------------ #
    # Public API
    # ------------------------------------------------------------------ #
    def to_dict(self) -> dict:
        """Return the full report content as a nested dict."""
        return {
            "title": self.settings.get("title", "Meta-analysis Report"),
            "authors": self.settings.get("authors", ""),
            "summary": self.meta_result.summary_text(),
            "characteristics_of_studies": self._characteristics_table(),
            "summary_of_findings": self._summary_of_findings(),
            "grade": self.settings.get("grade", None),
            "sensitivity": self._sensitivity_summary(),
            "subgroup": self._subgroup_summary(),
            "pooled_effect": self.meta_result.pooled_effect.to_dict(),
            "heterogeneity": self.meta_result.heterogeneity.to_dict(),
        }

    def to_markdown(self) -> str:
        """Generate the full PRISMA + Cochrane-style Markdown report."""
        sections: List[str] = []
        title = self.settings.get("title", "Meta-analysis Report")
        authors = self.settings.get("authors", "")
        sections.append(f"# {title}\n")
        if authors:
            sections.append(f"**Authors:** {authors}\n")
        # Summary.
        sections.append("## Summary\n")
        sections.append(self.meta_result.summary_text() + "\n")
        sections.append("\n### Pooled effect\n")
        pe = self.meta_result.pooled_effect
        ci_str = (
            f"**{pe.value:.3f}** (95% CI {pe.ci_lower:.3f} to {pe.ci_upper:.3f})"
            if pe.ci_lower is not None and pe.ci_upper is not None
            else f"**{pe.value:.3f}**"
        )
        sections.append(
            f"Pooled {pe.type.value} = {ci_str} "
            f"(z = {self.meta_result.test_statistic:.3f}, "
            f"p = {self.meta_result.p_value:.4g}).\n"
        )
        # Heterogeneity.
        sections.append("\n### Heterogeneity\n")
        het = self.meta_result.heterogeneity
        sections.append(
            f"I² = {het.I_squared:.1f}%, τ² = {het.tau_squared:.4f}, "
            f"Q = {het.Q:.2f} (df = {het.df}, p = {het.p_value:.4g}).\n"
        )
        sections.append(f"_{het.interpretation}_\n")
        # Characteristics of included studies.
        sections.append("\n## Characteristics of included studies\n")
        sections.append(self._markdown_characteristics())
        # Forest plot reference (image saved alongside).
        sections.append("\n## Forest plot\n")
        sections.append(
            "![Forest plot](forest_plot.png)\n"
            "*Forest plot showing per-study effects (squares sized by weight) "
            "and the pooled effect (diamond).*\n"
        )
        # Funnel plot reference.
        sections.append("\n## Funnel plot & publication bias\n")
        fp = FunnelPlot(self.effect_sizes, pooled=self.meta_result.pooled_effect)
        egger_t, egger_p, egger_b = fp.eggers_test()
        begg_tau, begg_p = fp.beggs_test()
        fsn = fp.rosenthal_fail_safe_n()
        sections.append(
            "![Funnel plot](funnel_plot.png)\n\n"
            f"- **Egger's test**: t = {egger_t:.3f}, p = {egger_p:.4g} "
            f"(bias = {egger_b:.3f}).\n"
            f"- **Begg's test**: Kendall τ = {begg_tau:.3f}, p = {begg_p:.4g}.\n"
            f"- **Rosenthal fail-safe N**: {fsn}.\n"
        )
        # Sensitivity.
        sections.append("\n## Sensitivity analyses\n")
        sections.append(self._markdown_sensitivity())
        # Subgroup.
        if self.settings.get("subgroups"):
            sections.append("\n## Subgroup analysis\n")
            sections.append(self._markdown_subgroup())
        # SoF.
        sections.append("\n## Summary of findings (SoF)\n")
        sections.append(self._markdown_sof())
        # GRADE.
        if self.settings.get("grade"):
            sections.append("\n## GRADE assessment\n")
            sections.append(self._markdown_grade())
        return "\n".join(sections)

    def generate(self, output_path: str, format: str = "pdf") -> str:
        """Generate the report and write it to ``output_path``.

        Args:
            output_path: Destination file path.
            format: ``'pdf'`` | ``'docx'`` | ``'html'`` | ``'md'``.

        Returns:
            The output_path.
        """
        fmt = format.lower()
        if fmt == "md":
            with open(output_path, "w", encoding="utf-8") as fh:
                fh.write(self.to_markdown())
        elif fmt == "html":
            html = self._to_html()
            with open(output_path, "w", encoding="utf-8") as fh:
                fh.write(html)
        elif fmt == "pdf":
            self._to_pdf(output_path)
        elif fmt == "docx":
            self._to_docx(output_path)
        else:
            raise ValueError(f"Unsupported report format {format!r}.")
        # Also save the forest & funnel plots next to the report.
        try:
            import warnings  # type: ignore
            base_dir = os.path.dirname(os.path.abspath(output_path))
            fp = ForestPlot(
                effect_sizes=self.effect_sizes,
                pooled=self.meta_result.pooled_effect,
                title=self.settings.get("title", ""),
                weights=self.meta_result.weights,
            )
            fp.add_heterogeneity(
                f"I² = {self.meta_result.I_squared:.1f}%  "
                f"Q = {self.meta_result.Q_statistic:.2f}  "
                f"p = {self.meta_result.Q_p_value:.4g}"
            )
            fig = fp.render(style=self.settings.get("style", "cochrane"))
            with warnings.catch_warnings():
                warnings.simplefilter("ignore", UserWarning)
                fp.save(os.path.join(base_dir, "forest_plot.png"), format="png")
            # Funnel.
            funnel = FunnelPlot(self.effect_sizes, pooled=self.meta_result.pooled_effect)
            funnel.render()
            with warnings.catch_warnings():
                warnings.simplefilter("ignore", UserWarning)
                funnel.save(os.path.join(base_dir, "funnel_plot.png"), format="png")
        except Exception as e:
            logger.warning("Failed to save forest/funnel PNGs alongside report: %s", e)
        return output_path

    # ------------------------------------------------------------------ #
    # Markdown section builders
    # ------------------------------------------------------------------ #
    def _characteristics_table(self) -> List[dict]:
        """Build the characteristics-of-studies table data (list of dicts)."""
        rows: List[dict] = []
        for i, es in enumerate(self.effect_sizes):
            sid = es.study_id or f"study_{i}"
            meta = next((d for d in (self.study_data or []) if d.get("study_id") == sid), {})
            rows.append({
                "study_id": sid,
                "study_name": es.study_name or meta.get("study_name", sid),
                "year": meta.get("year") or es.year or "—",
                "authors": meta.get("authors", "—"),
                "population": meta.get("population", "—"),
                "intervention": es.group_intervention or meta.get("intervention", "—"),
                "comparator": es.group_control or meta.get("comparator", "—"),
                "outcome": meta.get("outcome", "—"),
                "design": meta.get("design", "RCT"),
                "n_total": es.n_total or meta.get("n_total", "—"),
                "effect_value": float(es.value),
                "ci_lower": float(es.ci_lower) if es.ci_lower is not None else None,
                "ci_upper": float(es.ci_upper) if es.ci_upper is not None else None,
            })
        return rows

    def _markdown_characteristics(self) -> str:
        rows = self._characteristics_table()
        header = ("| Study | Year | Authors | Population | Intervention | "
                  "Comparator | Outcome | Design | N | Effect (95% CI) |\n"
                  "|---|---|---|---|---|---|---|---|---|---|\n")
        body = []
        for r in rows:
            ci = (
                f"{r['effect_value']:.2f} ({r['ci_lower']:.2f}, {r['ci_upper']:.2f})"
                if r["ci_lower"] is not None and r["ci_upper"] is not None
                else f"{r['effect_value']:.2f}"
            )
            body.append(
                f"| {r['study_name']} | {r['year']} | {r['authors']} | "
                f"{r['population']} | {r['intervention']} | {r['comparator']} | "
                f"{r['outcome']} | {r['design']} | {r['n_total']} | {ci} |"
            )
        return header + "\n".join(body) + "\n"

    def _sensitivity_summary(self) -> dict:
        try:
            loo = SensitivityAnalysis.leave_one_out(self.effect_sizes, method="DL")
            influence = SensitivityAnalysis.influence_diagnosis(self.effect_sizes, method="DL")
        except Exception as e:
            logger.warning("Sensitivity analysis failed: %s", e)
            return {"leave_one_out": [], "influence": {}}
        loo_summary = []
        for i, r in enumerate(loo):
            if r is None:
                continue
            sid = self.effect_sizes[i].study_id or f"study_{i}"
            loo_summary.append({
                "omit": sid,
                "pooled": r.pooled_effect.value,
                "ci_lower": r.pooled_effect.ci_lower,
                "ci_upper": r.pooled_effect.ci_upper,
                "i_squared": r.I_squared,
            })
        return {"leave_one_out": loo_summary, "influence": influence}

    def _markdown_sensitivity(self) -> str:
        s = self._sensitivity_summary()
        out = ["### Leave-one-out\n",
               "| Omitted study | Pooled | 95% CI | I² |\n",
               "|---|---|---|---|\n"]
        for row in s.get("leave_one_out", []):
            out.append(
                f"| {row['omit']} | {row['pooled']:.3f} | "
                f"({row['ci_lower']:.3f}, {row['ci_upper']:.3f}) | "
                f"{row['i_squared']:.1f}% |"
            )
        out.append("\n### Influence diagnostics\n")
        out.append("| Study | DFFITS | Cook's D |\n|---|---|---|\n")
        for sid, diag in s.get("influence", {}).items():
            out.append(f"| {sid} | {diag['DFFITS']:.4f} | {diag['cooks_d']:.4f} |")
        return "\n".join(out) + "\n"

    def _subgroup_summary(self) -> dict:
        sg_map = self.settings.get("subgroups") or {}
        if not sg_map:
            return {}
        try:
            res = SubgroupAnalysis.analyze(self.effect_sizes, sg_map, method="DL")
        except Exception as e:
            logger.warning("Subgroup analysis failed: %s", e)
            return {}
        return res.to_dict()

    def _markdown_subgroup(self) -> str:
        sg_map = self.settings.get("subgroups") or {}
        if not sg_map:
            return "_No subgroup variable provided._\n"
        try:
            res = SubgroupAnalysis.analyze(self.effect_sizes, sg_map, method="DL")
        except Exception as e:
            return f"_Subgroup analysis failed: {e}_\n"
        return res.to_markdown() + "\n"

    def _summary_of_findings(self) -> dict:
        """Build the summary-of-findings (SoF) table data."""
        pe = self.meta_result.pooled_effect
        # Assume dichotomous if we have cell counts; compute NNT if RD.
        sof = {
            "outcome": self.settings.get("outcome", pe.type.value),
            "n_studies": self.meta_result.studies_count,
            "n_participants": self.meta_result.total_participants,
            "effect_type": pe.type.value,
            "pooled": float(pe.value),
            "ci_lower": float(pe.ci_lower) if pe.ci_lower is not None else None,
            "ci_upper": float(pe.ci_upper) if pe.ci_upper is not None else None,
            "i_squared": float(self.meta_result.I_squared),
            "p_value": float(self.meta_result.p_value),
            "method": self.meta_result.method.value,
        }
        # Quality assessment (GRADE) if provided.
        if self.settings.get("grade"):
            sof["grade_quality"] = self.settings["grade"].get("quality", "—")
        return sof

    def _markdown_sof(self) -> str:
        s = self._summary_of_findings()
        ci_str = (
            f"{s['ci_lower']:.3f} to {s['ci_upper']:.3f}"
            if s["ci_lower"] is not None and s["ci_upper"] is not None
            else "—"
        )
        grade_str = s.get("grade_quality", "—")
        out = (
            "| Outcome | Studies | Participants | Pooled effect (95% CI) | "
            "I² | p-value | Method | Quality |\n"
            "|---|---|---|---|---|---|---|---|\n"
            f"| {s['outcome']} | {s['n_studies']} | {s['n_participants']} | "
            f"{s['pooled']:.3f} ({ci_str}) | {s['i_squared']:.1f}% | "
            f"{s['p_value']:.4g} | {s['method']} | {grade_str} |\n"
        )
        return out

    def _markdown_grade(self) -> str:
        g = self.settings.get("grade", {})
        if not g:
            return "_No GRADE ratings provided._\n"
        out = ["| Domain | Rating | Comment |\n|---|---|---|\n"]
        for domain, value in g.items():
            if isinstance(value, dict):
                rating = value.get("rating", "—")
                comment = value.get("comment", "")
            else:
                rating = value
                comment = ""
            out.append(f"| {domain} | {rating} | {comment} |")
        return "\n".join(out) + "\n"

    # ------------------------------------------------------------------ #
    # HTML / PDF / DOCX output
    # ------------------------------------------------------------------ #
    def _to_html(self) -> str:
        """Render the Markdown report to HTML."""
        md = self.to_markdown()
        try:
            import markdown  # type: ignore
            html_body = markdown.markdown(md, extensions=["tables", "fenced_code"])
        except ImportError:
            # Minimal fallback: just wrap in <pre>.
            html_body = f"<pre>{md}</pre>"
        # Embed images as base64 if they exist.
        base_dir = os.path.dirname(os.path.abspath(__file__))
        for img_name in ["forest_plot.png", "funnel_plot.png"]:
            img_path = os.path.join(base_dir, img_name)
            if os.path.exists(img_path):
                with open(img_path, "rb") as fh:
                    b64 = base64.b64encode(fh.read()).decode("ascii")
                html_body = html_body.replace(
                    f"src=\"{img_name}\"",
                    f"src=\"data:image/png;base64,{b64}\"",
                )
        return (
            "<!DOCTYPE html>\n<html lang=\"en\">\n<head>\n"
            "<meta charset=\"utf-8\">\n"
            "<title>Meta-analysis Report</title>\n"
            "<style>\n"
            "body { font-family: 'Noto Sans SC', 'DejaVu Sans', Arial, sans-serif; "
            "max-width: 900px; margin: 2em auto; padding: 0 1em; color: #222; }\n"
            "table { border-collapse: collapse; width: 100%; margin: 1em 0; }\n"
            "th, td { border: 1px solid #ccc; padding: 6px 10px; text-align: left; }\n"
            "th { background: #f0f0f0; }\n"
            "img { max-width: 100%; height: auto; display: block; margin: 1em auto; }\n"
            "</style>\n</head>\n<body>\n"
            f"{html_body}\n</body>\n</html>"
        )

    def _to_pdf(self, output_path: str) -> None:
        """Render the report to PDF using reportlab (Markdown → PDF)."""
        try:
            from reportlab.lib.pagesizes import A4  # type: ignore
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image as RLImage,
            )  # type: ignore
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError(
                "reportlab is required to render PDF reports; "
                "install with `pip install reportlab`."
            ) from exc
        md = self.to_markdown()
        doc = SimpleDocTemplate(output_path, pagesize=A4,
                               topMargin=40, bottomMargin=40,
                               leftMargin=50, rightMargin=50)
        styles = getSampleStyleSheet()
        body_style = styles["BodyText"]
        h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=18,
                             spaceAfter=12, fontName="Helvetica-Bold")
        h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=14,
                             spaceAfter=8, fontName="Helvetica-Bold")
        h3 = ParagraphStyle("H3", parent=styles["Heading3"], fontSize=12,
                             spaceAfter=6, fontName="Helvetica-Bold")
        flow = []
        for line in md.split("\n"):
            stripped = line.strip()
            if not stripped:
                flow.append(Spacer(1, 6))
                continue
            if stripped.startswith("# "):
                flow.append(Paragraph(stripped[2:], h1))
            elif stripped.startswith("## "):
                flow.append(Paragraph(stripped[3:], h2))
            elif stripped.startswith("### "):
                flow.append(Paragraph(stripped[4:], h3))
            elif stripped.startswith("|"):
                # Render table rows as plain paragraphs (limited table support).
                flow.append(Paragraph(stripped.replace("|", " | "), body_style))
            elif stripped.startswith("!["):
                # Inline image.
                img_match = stripped.split("](")[1].rstrip(")")
                img_path = os.path.join(
                    os.path.dirname(os.path.abspath(output_path)), img_match
                )
                if os.path.exists(img_path):
                    flow.append(RLImage(img_path, width=460, height=360))
            else:
                # Strip Markdown italics/bold for PDF rendering.
                clean = stripped.replace("*", "").replace("_", "")
                flow.append(Paragraph(clean, body_style))
        doc.build(flow)

    def _to_docx(self, output_path: str) -> None:
        """Render the report to a .docx file."""
        try:
            import docx  # type: ignore
            from docx.shared import Inches  # type: ignore
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError(
                "python-docx is required to render DOCX reports; "
                "install with `pip install python-docx`."
            ) from exc
        doc = docx.Document()
        title = self.settings.get("title", "Meta-analysis Report")
        doc.add_heading(title, level=0)
        # Summary.
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(self.meta_result.summary_text())
        # Pooled.
        pe = self.meta_result.pooled_effect
        ci_str = (
            f" (95% CI {pe.ci_lower:.3f} to {pe.ci_upper:.3f})"
            if pe.ci_lower is not None and pe.ci_upper is not None else ""
        )
        doc.add_paragraph(
            f"Pooled {pe.type.value} = {pe.value:.3f}{ci_str} "
            f"(z = {self.meta_result.test_statistic:.3f}, "
            f"p = {self.meta_result.p_value:.4g})."
        )
        # Heterogeneity.
        doc.add_heading("Heterogeneity", level=2)
        het = self.meta_result.heterogeneity
        doc.add_paragraph(
            f"I² = {het.I_squared:.1f}%, τ² = {het.tau_squared:.4f}, "
            f"Q = {het.Q:.2f} (df = {het.df}, p = {het.p_value:.4g}). "
            f"{het.interpretation}"
        )
        # Characteristics table.
        doc.add_heading("Characteristics of included studies", level=1)
        rows = self._characteristics_table()
        if rows:
            tbl = doc.add_table(rows=1 + len(rows), cols=5)
            tbl.style = "Light List Accent 1"
            hdr = tbl.rows[0].cells
            for i, col in enumerate(["Study", "Year", "Intervention", "N", "Effect (95% CI)"]):
                hdr[i].text = col
            for ri, r in enumerate(rows, start=1):
                row_cells = tbl.rows[ri].cells
                row_cells[0].text = str(r["study_name"])
                row_cells[1].text = str(r["year"])
                row_cells[2].text = str(r["intervention"])
                row_cells[3].text = str(r["n_total"])
                ci = (
                    f"{r['effect_value']:.2f} ({r['ci_lower']:.2f}, {r['ci_upper']:.2f})"
                    if r["ci_lower"] is not None and r["ci_upper"] is not None
                    else f"{r['effect_value']:.2f}"
                )
                row_cells[4].text = ci
        # Embed forest plot.
        doc.add_heading("Forest plot", level=1)
        try:
            base_dir = os.path.dirname(os.path.abspath(output_path))
            forest_path = os.path.join(base_dir, "forest_plot.png")
            if os.path.exists(forest_path):
                doc.add_picture(forest_path, width=Inches(6.0))
        except Exception as e:
            logger.warning("Failed to embed forest plot in DOCX: %s", e)
        # Embed funnel plot.
        doc.add_heading("Funnel plot", level=1)
        try:
            base_dir = os.path.dirname(os.path.abspath(output_path))
            funnel_path = os.path.join(base_dir, "funnel_plot.png")
            if os.path.exists(funnel_path):
                doc.add_picture(funnel_path, width=Inches(5.0))
        except Exception as e:
            logger.warning("Failed to embed funnel plot in DOCX: %s", e)
        # Summary of findings.
        doc.add_heading("Summary of findings", level=1)
        sof = self._summary_of_findings()
        ci = (
            f"{sof['ci_lower']:.3f} to {sof['ci_upper']:.3f}"
            if sof["ci_lower"] is not None and sof["ci_upper"] is not None else "—"
        )
        doc.add_paragraph(
            f"Outcome: {sof['outcome']}; Studies: {sof['n_studies']}; "
            f"Participants: {sof['n_participants']}; "
            f"Pooled effect: {sof['pooled']:.3f} (95% CI {ci}); "
            f"I²: {sof['i_squared']:.1f}%; p: {sof['p_value']:.4g}; "
            f"Method: {sof['method']}."
        )
        doc.save(output_path)
