"""Word (.docx) report generator built on ``python-docx``.

This module produces styled Word documents with headings, paragraphs,
page breaks, an auto-updating Table of Contents, embedded charts, and a
formatted bibliography section in APA / IEEE / MLA style.

Heavy deps (``python-docx``, ``matplotlib``, ``Pillow``) are imported
lazily inside the methods so the module itself is always importable.
"""
#
# MIT License — Academic Research Suite
# Copyright (c) 2026 — see /LICENSE for full text.
#
from __future__ import annotations

import io
import logging
import os
import re
from typing import Any, Dict, List, Optional, Sequence, Union

from ._paper_utils import (
    get_authors,
    get_citation_count,
    get_doi,
    get_field,
    get_keywords,
    get_str,
    get_year,
)

logger = logging.getLogger(__name__)


# Default paper-table columns for DOCX export.
DEFAULT_DOCX_COLUMNS: List[str] = [
    "title", "authors", "year", "journal", "citations",
]


# ---------------------------------------------------------------------------
# Bibliography formatters: APA, IEEE, MLA.
# ---------------------------------------------------------------------------
def _format_authors_apa(authors: List[str]) -> str:
    """Format authors in APA style: 'Last, F. M., & Last2, F. M.'."""
    out: List[str] = []
    for a in authors:
        a = a.strip().rstrip(".")
        if "," in a:
            last, rest = a.split(",", 1)
            last = last.strip()
            rest = rest.strip()
            initials = " ".join(f"{p[0]}." for p in rest.split() if p)
            out.append(f"{last}, {initials}")
        else:
            parts = a.split()
            if len(parts) >= 2:
                last = parts[-1]
                initials = " ".join(f"{p[0]}." for p in parts[:-1] if p)
                out.append(f"{last}, {initials}")
            else:
                out.append(a)
    if len(out) == 1:
        return out[0]
    if len(out) == 2:
        return f"{out[0]}, & {out[1]}"
    return ", ".join(out[:-1]) + ", & " + out[-1]


def _format_authors_ieee(authors: List[str]) -> str:
    """Format authors in IEEE style: 'F. M. Last, F. M. Last2, and F. M. Last3'."""
    out: List[str] = []
    for a in authors:
        a = a.strip().rstrip(".")
        if "," in a:
            last, rest = a.split(",", 1)
            last = last.strip()
            rest = rest.strip()
            initials = " ".join(f"{p[0]}." for p in rest.split() if p)
            out.append(f"{initials} {last}")
        else:
            parts = a.split()
            if len(parts) >= 2:
                last = parts[-1]
                initials = " ".join(f"{p[0]}." for p in parts[:-1] if p)
                out.append(f"{initials} {last}")
            else:
                out.append(a)
    if len(out) == 1:
        return out[0]
    if len(out) <= 6:
        return ", ".join(out[:-1]) + ", and " + out[-1]
    return ", ".join(out[:6]) + ", et al."


def _format_authors_mla(authors: List[str]) -> str:
    """Format authors in MLA style: 'Last, First M., and First M. Last2.'."""
    if not authors:
        return ""
    if len(authors) == 1:
        return authors[0].strip().rstrip(".")
    if len(authors) == 2:
        # MLA: "Last, First, and First Last".
        a1 = authors[0].strip().rstrip(".")
        a2 = authors[1].strip().rstrip(".")
        if "," in a2:
            parts = a2.split(",", 1)
            a2 = f"{parts[1].strip()} {parts[0].strip()}"
        return f"{a1}, and {a2}"
    # 3+ authors: MLA lists only the first author + 'et al.'
    a1 = authors[0].strip().rstrip(".")
    return f"{a1}, et al."


def _format_paper_apa(paper: Any) -> str:
    """Render a single paper as an APA 7th reference entry."""
    authors = get_authors(paper)
    year = get_year(paper) or "n.d."
    title = get_str(paper, "title")
    journal = get_str(paper, "journal") or get_str(paper, "booktitle")
    volume = get_str(paper, "volume")
    number = get_str(paper, "number")
    pages = get_str(paper, "pages")
    doi = get_doi(paper)
    parts: List[str] = []
    if authors:
        parts.append(_format_authors_apa(authors) + ".")
    parts.append(f"({year}).")
    if title:
        parts.append(f"{title}.")
    if journal:
        parts.append(f"<i>{journal}</i>")
        if volume:
            vol_str = f", <b>{volume}</b>"
            if number:
                vol_str += f"({number})"
            parts.append(vol_str)
        if pages:
            parts.append(f", {pages}.")
        else:
            parts.append(".")
    else:
        parts.append(".")
    if doi:
        parts.append(f"https://doi.org/{doi}")
    return " ".join(p.strip() for p in parts if p.strip())


def _format_paper_ieee(paper: Any) -> str:
    """Render a single paper as an IEEE-style reference entry."""
    authors = get_authors(paper)
    year = get_year(paper) or "n.d."
    title = get_str(paper, "title")
    journal = get_str(paper, "journal") or get_str(paper, "booktitle")
    volume = get_str(paper, "volume")
    number = get_str(paper, "number")
    pages = get_str(paper, "pages")
    doi = get_doi(paper)
    parts: List[str] = []
    if authors:
        parts.append(_format_authors_ieee(authors) + ",")
    if title:
        # IEEE uses double quotes around article titles.
        parts.append(f'"{title},"')
    if journal:
        italic_parts: List[str] = [f"<i>{journal}</i>"]
        if volume:
            italic_parts.append(f"vol. {volume}")
        if number:
            italic_parts.append(f"no. {number}")
        if pages:
            parts.append(", ".join(italic_parts) + f", pp. {pages}")
        else:
            parts.append(", ".join(italic_parts))
        parts.append(f", {year}.")
    else:
        parts.append(f"{year}.")
    if doi:
        parts.append(f"doi: {doi}.")
    return " ".join(p.strip() for p in parts if p.strip())


def _format_paper_mla(paper: Any) -> str:
    """Render a single paper as an MLA 9th-style reference entry."""
    authors = get_authors(paper)
    title = get_str(paper, "title")
    journal = get_str(paper, "journal") or get_str(paper, "booktitle")
    volume = get_str(paper, "volume")
    number = get_str(paper, "number")
    year = get_year(paper) or "n.d."
    pages = get_str(paper, "pages")
    doi = get_doi(paper)
    parts: List[str] = []
    if authors:
        parts.append(_format_authors_mla(authors) + ".")
    if title:
        parts.append(f'"{title}."')
    if journal:
        parts.append(f"<i>{journal}</i>,")
        container_bits: List[str] = []
        if volume:
            container_bits.append(f"vol. {volume}")
        if number:
            container_bits.append(f"no. {number}")
        if container_bits:
            parts.append(", ".join(container_bits) + ",")
        parts.append(f"{year},")
        if pages:
            parts.append(f"pp. {pages}.")
        else:
            parts.append(".")
    else:
        parts.append(f"{year}.")
    if doi:
        parts.append(f"https://doi.org/{doi}.")
    return " ".join(p.strip() for p in parts if p.strip())


_BIB_FORMATTERS: Dict[str, Any] = {
    "apa": _format_paper_apa,
    "ieee": _format_paper_ieee,
    "mla": _format_paper_mla,
}


# ---------------------------------------------------------------------------
# Inline-markup parser: convert "<i>..</i>" / "<b>..</b>" to docx runs.
# ---------------------------------------------------------------------------
_INLINE_TAG_RE = re.compile(
    r"<(/?)(b|i|strong|em|u|sub|sup)\s*>",
    re.IGNORECASE,
)


def _add_runs_with_markup(paragraph: Any, text: str) -> None:
    """Add runs to ``paragraph`` honouring ``<b>``, ``<i>`` etc. tags.

    Recognised tags: ``<b>``, ``<strong>``, ``<i>``, ``<em>``, ``<u>``,
    ``<sub>``, ``<sup>``.  Unknown tags are stripped (text kept).
    """
    if not text:
        return
    # Tokenise: literal text segments + tag tokens.
    pos = 0
    bold = italic = underline = subscript = superscript = False
    for m in _INLINE_TAG_RE.finditer(text):
        if m.start() > pos:
            chunk = text[pos:m.start()]
            if chunk:
                paragraph.add_run(chunk)
                _apply_run_style(paragraph.runs[-1], bold, italic, underline, subscript, superscript)
        pos = m.end()
        closing, name = m.group(1), m.group(2).lower()
        if name in ("b", "strong"):
            bold = (closing == "")
        elif name in ("i", "em"):
            italic = (closing == "")
        elif name == "u":
            underline = (closing == "")
        elif name == "sub":
            subscript = (closing == "")
        elif name == "sup":
            superscript = (closing == "")
    if pos < len(text):
        chunk = text[pos:]
        if chunk:
            paragraph.add_run(chunk)
            _apply_run_style(paragraph.runs[-1], bold, italic, underline, subscript, superscript)


def _apply_run_style(run: Any, bold: bool, italic: bool, underline: bool,
                     subscript: bool, superscript: bool) -> None:
    """Mutate the last run of a paragraph with the active style flags.

    python-docx requires setting flags AFTER the run is created, hence
    this helper.
    """
    try:
        run.bold = run.bold or bold
        run.italic = run.italic or italic
        run.underline = run.underline or underline
        if subscript:
            run.font.subscript = True
        if superscript:
            run.font.superscript = True
    except Exception as exc:  # pragma: no cover
        logger.debug("run style apply failed: %s", exc)


# ---------------------------------------------------------------------------
# DOCXReport class.
# ---------------------------------------------------------------------------
class DOCXReport:
    """Build a Word document report with TOC + bibliography."""

    def __init__(self, title: str, author: Optional[str] = None) -> None:
        """Initialise the report.

        Args:
            title: Document title (used for core properties + heading).
            author: Optional author for core properties.
        """
        try:
            from docx import Document  # noqa: WPS433
            from docx.shared import Pt, RGBColor, Inches  # noqa: WPS433
        except ImportError as exc:  # pragma: no cover
            logger.error("python-docx not installed: %s", exc)
            raise
        self.title = title
        self.author = author
        self._doc = Document()
        # Core properties.
        try:
            cp = self._doc.core_properties
            cp.title = title
            if author:
                cp.author = author
        except Exception as exc:  # pragma: no cover
            logger.debug("core_properties set failed: %s", exc)
        # Default font: try to set a CJK-friendly fallback via styles.
        try:
            normal = self._doc.styles["Normal"]
            normal.font.name = "Calibri"
            normal.font.size = Pt(11)
            # Asian font fallback (Calibri does not cover CJK; Word will
            # substitute its default Asian font, typically SimSun / MS Gothic,
            # which is fine for cross-platform display).
        except Exception as exc:  # pragma: no cover
            logger.debug("Normal style setup failed: %s", exc)
        # Title heading.
        try:
            h = self._doc.add_heading(title, level=0)
            h.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
        except Exception as exc:  # pragma: no cover
            logger.debug("title heading failed: %s", exc)

    # ------------------------------------------------------------------
    # Content API
    # ------------------------------------------------------------------
    def add_heading(self, text: str, level: int = 1) -> None:
        """Append a heading paragraph at the given outline level.

        Args:
            text: Heading text (may contain ``<b>``, ``<i>`` tags).
            level: 0 for the document title (Title style), 1-9 otherwise.
        """
        try:
            para = self._doc.add_heading(level=level)
            _add_runs_with_markup(para, text)
        except Exception as exc:
            logger.warning("add_heading failed: %s", exc, exc_info=True)
            self._doc.add_paragraph(text)

    def add_paragraph(self, text: str, style: Optional[str] = None) -> None:
        """Append a paragraph with optional style and inline markup.

        Args:
            text: Paragraph text. May contain ``<b>``, ``<i>``, ``<u>``,
                ``<sub>``, ``<sup>`` inline tags.
            style: Optional Word style name (e.g. ``'List Bullet'``).
        """
        try:
            para = self._doc.add_paragraph(style=style) if style else self._doc.add_paragraph()
            _add_runs_with_markup(para, text)
        except Exception as exc:
            logger.warning("add_paragraph failed: %s", exc, exc_info=True)
            self._doc.add_paragraph(text)

    def add_page_break(self) -> None:
        """Insert an explicit page break."""
        try:
            self._doc.add_page_break()
        except Exception as exc:  # pragma: no cover
            logger.debug("add_page_break failed: %s", exc)

    # ------------------------------------------------------------------
    # Papers table
    # ------------------------------------------------------------------
    def add_papers_table(
        self,
        papers: Sequence[Any],
        columns: Optional[List[str]] = None,
    ) -> None:
        """Append a styled Word table of papers with a shaded header row.

        Args:
            papers: Sequence of Paper-like objects.
            columns: Optional override list of column names.  Defaults to
                :data:`DEFAULT_DOCX_COLUMNS`.
        """
        cols = list(columns) if columns else list(DEFAULT_DOCX_COLUMNS)
        n_rows = len(papers) + 1  # +1 for header.
        table = self._doc.add_table(rows=n_rows, cols=len(cols))
        table.style = "Light Grid Accent 1"
        try:
            table.autofit = True
        except Exception:  # pragma: no cover
            pass

        # Header row.
        for j, col in enumerate(cols):
            cell = table.cell(0, j)
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(col.replace("_", " ").title())
            run.bold = True
            try:
                from docx.shared import RGBColor  # noqa: WPS433
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                from docx.oxml.ns import nsdecls  # noqa: WPS433
                from docx.oxml import parse_xml  # noqa: WPS433
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="2E5C8A" w:val="clear"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading)
            except Exception as exc:  # pragma: no cover
                logger.debug("header shading failed: %s", exc)

        # Body rows.
        for i, paper in enumerate(papers, start=1):
            for j, col in enumerate(cols):
                cell = table.cell(i, j)
                cell.text = ""
                para = cell.paragraphs[0]
                _add_runs_with_markup(para, self._cell_value(paper, col))

    def _cell_value(self, paper: Any, col: str) -> str:
        """Return a short string representation of one cell."""
        if col == "authors":
            authors = get_authors(paper)
            if not authors:
                return ""
            if len(authors) <= 3:
                return "; ".join(authors)
            return f"{authors[0]} <i>et al.</i> ({len(authors)})"
        if col == "citations":
            return str(get_citation_count(paper))
        if col == "year":
            y = get_year(paper)
            return str(y) if y is not None else ""
        if col == "doi":
            return get_doi(paper)
        if col == "title":
            t = get_str(paper, "title")
            return t if len(t) <= 120 else t[:117] + "..."
        return get_str(paper, col)

    # ------------------------------------------------------------------
    # Chart image
    # ------------------------------------------------------------------
    def add_chart(
        self,
        figure_or_path: Union[Any, str, bytes],
        caption: Optional[str] = None,
        width_inches: float = 6,
    ) -> None:
        """Embed a chart image (Figure object, path, or raw bytes).

        Args:
            figure_or_path: ``matplotlib.figure.Figure`` / file path / raw
                image bytes.
            caption: Optional caption text (centred, italic).
            width_inches: Image width in inches.
        """
        from docx.shared import Inches  # noqa: WPS433
        img_path = self._materialise_image(figure_or_path)
        if img_path is None:
            return
        try:
            para = self._doc.add_paragraph()
            para.alignment = 1  # CENTER
            run = para.add_run()
            run.add_picture(img_path, width=Inches(width_inches))
            if caption:
                cap_para = self._doc.add_paragraph(caption)
                cap_para.alignment = 1
                for r in cap_para.runs:
                    r.italic = True
                    try:
                        from docx.shared import RGBColor  # noqa: WPS433
                        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    except Exception:  # pragma: no cover
                        pass
        except Exception as exc:
            logger.error("add_chart failed: %s", exc, exc_info=True)
        finally:
            # Clean up temporary files we created.
            if isinstance(figure_or_path, (bytes, bytearray)) and img_path:
                try:
                    os.unlink(img_path)
                except OSError:
                    pass

    def _materialise_image(self, figure_or_path: Union[Any, str, bytes]) -> Optional[str]:
        """Return a filesystem path to a PNG/JP2 image suitable for docx."""
        if isinstance(figure_or_path, str):
            return figure_or_path if os.path.exists(figure_or_path) else None
        if isinstance(figure_or_path, (bytes, bytearray)):
            fd, tmp = __import__("tempfile").mkstemp(suffix=".png")
            try:
                with os.fdopen(fd, "wb") as fh:
                    fh.write(figure_or_path)
            except Exception as exc:  # pragma: no cover
                logger.error("temp image write failed: %s", exc)
                return None
            return tmp
        # matplotlib Figure
        try:
            fd, tmp = __import__("tempfile").mkstemp(suffix=".png")
            with os.fdopen(fd, "wb") as fh:
                figure_or_path.savefig(fh, format="png", dpi=150)
            return tmp
        except Exception as exc:
            logger.error("figure savefig failed: %s", exc, exc_info=True)
            return None

    # ------------------------------------------------------------------
    # TOC (auto-updates in Word)
    # ------------------------------------------------------------------
    def add_toc(self) -> None:
        """Insert a Table of Contents field.

        The TOC is populated by Word/LibreOffice on first open or by pressing
        ``F9`` (Word) / ``Tools > Update > All`` (LibreOffice).  A fallback
        placeholder paragraph is also inserted so the field is visually
        recognisable before the first refresh.
        """
        try:
            from docx.oxml.ns import qn  # noqa: WPS433
            from docx.oxml import OxmlElement  # noqa: WPS433
        except ImportError as exc:  # pragma: no cover
            logger.warning("oxml unavailable; TOC cannot be inserted: %s", exc)
            return
        para = self._doc.add_paragraph()
        run = para.add_run()
        # Build the TOC field XML: begin, instrText, separate, end.
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        # \o "1-3" = include levels 1-3, \h = hyperlink, \z = hide page numbers
        # in web view, \u = use applied paragraph outline levels.
        instr.text = 'TOC \\o "1-3" \\h \\z \\u'
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        # Placeholder text shown until Word refreshes the field.
        placeholder = OxmlElement("w:r")
        placeholder_text = OxmlElement("w:t")
        placeholder_text.text = "Right-click here and choose 'Update Field' to populate the Table of Contents."
        placeholder.append(placeholder_text)
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_sep)
        run._r.append(placeholder)
        run._r.append(fld_end)
        # Add a page break after the TOC.
        self.add_page_break()

    # ------------------------------------------------------------------
    # Bibliography
    # ------------------------------------------------------------------
    def add_bibliography(
        self,
        papers: Sequence[Any],
        style: str = "apa",
    ) -> None:
        """Append a formatted bibliography section.

        Args:
            papers: Sequence of Paper-like objects.
            style: One of ``'apa'``, ``'ieee'``, ``'mla'``.
        """
        style = style.lower()
        if style not in _BIB_FORMATTERS:
            logger.warning("Unknown bibliography style %r; defaulting to 'apa'", style)
            style = "apa"
        formatter = _BIB_FORMATTERS[style]
        self.add_heading("Bibliography", level=1)
        # Sort by first author surname, then year.
        def _sort_key(p: Any) -> tuple:
            authors = get_authors(p)
            first = authors[0] if authors else ""
            if "," in first:
                surname = first.split(",")[0].strip().lower()
            else:
                surname = first.split()[-1].lower() if first else ""
            return (surname, get_year(p) or 0)
        sorted_papers = sorted(papers, key=_sort_key)
        for paper in sorted_papers:
            text = formatter(paper)
            try:
                para = self._doc.add_paragraph(style="List Number" if style == "ieee" else None)
                _add_runs_with_markup(para, text)
                # Hanging indent for APA / MLA.
                if style in {"apa", "mla"}:
                    try:
                        from docx.shared import Inches  # noqa: WPS433
                        para_format = para.paragraph_format
                        para_format.left_indent = Inches(0.5)
                        para_format.first_line_indent = Inches(-0.5)
                    except Exception as exc:  # pragma: no cover
                        logger.debug("bib indent failed: %s", exc)
            except Exception as exc:
                logger.error("bibliography entry failed: %s", exc, exc_info=True)

    # ------------------------------------------------------------------
    # Build
    # ------------------------------------------------------------------
    def build(self, path: str) -> str:
        """Save the document to ``path`` and return the absolute path.

        Args:
            path: Target ``.docx`` file path (parent dirs auto-created).

        Returns:
            The absolute path written.
        """
        os.makedirs(os.path.dirname(os.path.abspath(path)) or ".", exist_ok=True)
        abs_path = os.path.abspath(path)
        try:
            self._doc.save(abs_path)
        except Exception as exc:
            logger.error("DOCX save failed -> %s: %s", abs_path, exc, exc_info=True)
            raise
        logger.info("DOCX built -> %s", abs_path)
        return abs_path


__all__ = ["DOCXReport"]
